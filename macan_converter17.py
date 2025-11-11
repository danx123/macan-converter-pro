import sys
import os
import pypdfium2 as pdfium
import re
import subprocess
from PIL import Image
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QFileDialog, QProgressBar,
    QComboBox, QMessageBox, QFrame, QStackedWidget, QCheckBox, QGridLayout,
    QTextEdit, QListWidget, QListWidgetItem
)
from PySide6.QtCore import QThread, QObject, Signal, Qt, QSize, QByteArray
from PySide6.QtGui import QIcon, QPixmap

# [MODIFIKASI] Mengimpor library baru yang dibutuhkan dan memperbarui pesan error
try:
    import docx
    from docx2pdf import convert
    import openpyxl
except ImportError:
    print("PERINGATAN: Beberapa fungsionalitas mungkin tidak berfungsi. Pastikan Anda menginstal library yang dibutuhkan dengan 'pip install python-docx docx2pdf openpyxl pypdfium2'")


# --- Kamus Teks Multi-Bahasa (Tidak Berubah) ---
LANGUAGES = {
    "id": {
        # Judul & Tab
        "window_title": "Macan Converter Pro",
        "main_title": "Macan Converter Pro",
        "tab_file": "Dokumen",
        "tab_image": "Gambar",
        "tab_video": "Video",
        "tab_audio": "Audio",
        "tab_video_audio": "Ekstrak Audio", # Diubah untuk kejelasan
        "tab_about": "Tentang",

        # UI Umum
        "input_label": "Pilih File Input:",
        "output_label": "Pilih Folder Output:",
        "browse_btn": "Browse...",
        "start_conversion_btn": "Mulai Konversi",
        "stop_conversion_btn": "Hentikan",
        "output_folder_btn": "Folder Output",
        "ready_status": "Siap untuk mengonversi.",
        "batch_mode_checkbox": "Batch Mode (Konversi Banyak File Sekaligus)",
        "after_converting_label": "Setelah Konversi:",
        "shutdown_checkbox": "Matikan PC Ketika Selesai",

        # File Converter (Sekarang Dokumen)
        "input_placeholder_file": "Pilih file yang akan dikonversi...",
        "output_placeholder_folder": "Pilih folder untuk menyimpan hasil...",
        "conversion_mode_label": "Pilih Mode Konversi:",
        "mode_label": "Mode:",
        "output_format_label": "Format Output:",
        "ico_res_label": "Pilih Resolusi ICO:",
        "modes": ["PDF ke Gambar", "PNG ke ICO", "PDF ke TXT", "PDF ke DOCX", "PDF ke XLSX", "DOCX ke PDF", "Gambar ke PDF"],
        
        # Image Converter
        "input_img_label": "Pilih File Gambar Input:",
        "input_placeholder_single_img": "Pilih satu file gambar...",
        "input_placeholder_multi_img": "Pilih satu atau beberapa file gambar...",
        "output_placeholder_img": "Pilih folder tujuan...",
        "img_options_label": "Opsi Konversi Gambar:",
        "resolution_label": "Resolusi:",
        "quality_label": "Kualitas:",
        "img_formats": ["JPG", "PNG", "WEBP", "BMP", "GIF"],
        "img_qualities": ["Maksimum (100)", "Sangat Baik (95)", "Baik (85)", "Sedang (75)", "Rendah (50)"],
        
        # Video Converter
        "input_vid_label": "Pilih File Video Input:",
        "input_placeholder_single_vid": "Pilih satu file video...",
        "input_placeholder_multi_vid": "Pilih beberapa file video untuk konversi batch...",
        "output_placeholder_vid": "Pilih folder tujuan...",
        "vid_options_label": "Opsi Konversi Video:",
        "vid_formats": ["mp4", "mkv", "avi", "mov", "webm", "gif"],
        "vid_qualities": ["Tinggi", "Sedang", "Rendah"],
        "shutdown_confirm_title": "Konfirmasi Matikan PC",
        "shutdown_confirm_msg": "Konversi selesai. Apakah Anda yakin ingin mematikan PC sekarang?",

        # Audio Converter
        "input_audio_label": "Pilih File Audio Input:",
        "input_placeholder_single_audio": "Pilih satu file audio...",
        "input_placeholder_multi_audio": "Pilih beberapa file audio untuk konversi batch...",
        "output_placeholder_audio": "Pilih folder tujuan...",
        "audio_options_label": "Opsi Konversi Audio:",
        "bitrate_label": "Bitrate:",
        "audio_formats": ["mp3", "wav", "aac", "flac", "ogg", "wma", "m4a"],

        # Video-Audio Converter (Ekstrak Audio)
        "input_va_label": "Pilih File Video Input:",
        "output_va_label": "Pilih Folder Output Audio:",
        "input_placeholder_single_va": "Pilih satu file video...",
        "input_placeholder_multi_va": "Pilih beberapa video untuk konversi batch...",
        "output_placeholder_va": "Pilih folder tujuan audio...",
        "va_options_label": "Opsi Ekstraksi Audio:",
        
        # Pesan & Notifikasi
        "invalid_input_title": "Input Tidak Valid",
        "invalid_input_file_msg": "Silakan pilih file input yang valid.",
        "invalid_output_folder_msg": "Silakan pilih folder output yang valid.",
        "no_ico_size_msg": "Pilih setidaknya satu ukuran resolusi untuk file ICO.",
        "conversion_logic_error": "Logic konversi untuk mode '{mode}' tidak ditemukan.",
        "docx_conv_error_title": "Error Konversi DOCX",
        "docx_conv_error_msg": "Konversi DOCX ke PDF gagal. Pastikan Microsoft Word (Windows) atau LibreOffice (Linux/Mac) terinstal.",
        "ffmpeg_not_found_title": "FFmpeg Tidak Ditemukan",
        "ffmpeg_not_found_msg": "ffmpeg tidak ditemukan di direktori aplikasi atau di sistem PATH. Silakan letakkan file 'ffmpeg.exe' (Windows) atau 'ffmpeg' (Mac/Linux) di sebelah aplikasi ini.",
        "batch_no_files_msg": "Silakan pilih setidaknya satu file untuk mode batch.",
        "conversion_success_msg": "Sukses! Konversi telah selesai.",
        "batch_complete_title": "Batch Selesai",
        "batch_complete_msg": "Batch selesai! Semua {count} file telah berhasil dikonversi.",
        "error_title": "Error",
        "error_during_conversion": "Terjadi kesalahan: {error}",
        "error_in_batch": "Error pada file {index}: {error}\n\nProses batch dihentikan.",
        "conversion_stopped": "Konversi dihentikan oleh pengguna.",
        
        # Status Konversi
        "converting_page": "Mengonversi halaman {current} dari {total}...",
        "pdf_conversion_success": "Sukses! {total} halaman PDF telah dikonversi.",
        "opening_png": "Membuka file PNG...",
        "saving_ico": "Menyimpan ke format ICO dengan ukuran yang dipilih...",
        "done": "Selesai!",
        "png_to_ico_success": "Sukses! File PNG telah dikonversi ke ICO.",
        "opening_image": "Membuka {filename}...",
        "processing_image": "Memproses gambar...",
        "saving_image": "Gambar berhasil disimpan.",
        "image_conversion_success": "Sukses! Gambar telah dikonversi ke {format}.",
        "video_conversion_success": "Sukses! Video telah dikonversi ke {format}.",
        "getting_video_info": "Mendapatkan info video...",
        "converting_progress": "Mengonversi... {progress}%",
        "getting_audio_info": "Mendapatkan info audio...",
        "audio_conversion_success": "Sukses! Audio telah dikonversi ke {format}.",
        "preparing_conversion": "Mempersiapkan konversi untuk {filename}...",
        "converting_batch_file": "Mengonversi file {current} dari {total}: {filename}",
        
        # Tab About
        "about_content": """<b>Macan Converter Pro</b>
Versi 3.0 (Redesign)

Dikembangkan oleh: Macan Angkasa
© 2025 Danx Exodus

Macan Converter Pro adalah aplikasi multi-konversi modern
untuk file, gambar, audio, dan video.
Mendukung batch mode, berbagai format populer, serta antarmuka
bersih dan mudah digunakan."""
    },
    "en": {
        # Titles & Tabs
        "window_title": "Macan Converter Pro",
        "main_title": "Macan Converter Pro",
        "tab_file": "Document",
        "tab_image": "Picture",
        "tab_video": "Video",
        "tab_audio": "Audio",
        "tab_video_audio": "Extract Audio", # Changed for clarity
        "tab_about": "About",

        # Common UI
        "input_label": "Select Input File(s):",
        "output_label": "Select Output Folder:",
        "browse_btn": "Browse...",
        "start_conversion_btn": "Start",
        "stop_conversion_btn": "Stop",
        "output_folder_btn": "Output Folder",
        "ready_status": "Ready to convert.",
        "batch_mode_checkbox": "Batch Mode (Convert Multiple Files at Once)",
        "after_converting_label": "After Converting:",
        "shutdown_checkbox": "Shutdown when done",

        # File Converter (Now Document)
        "input_placeholder_file": "Select a file to convert...",
        "output_placeholder_folder": "Select a folder to save the result...",
        "conversion_mode_label": "Select Conversion Mode:",
        "mode_label": "Mode:",
        "output_format_label": "Output Format:",
        "ico_res_label": "Select ICO Resolutions:",
        "modes": ["PDF to Image", "PNG to ICO", "PDF to TXT", "PDF to DOCX", "PDF to XLSX", "DOCX to PDF", "Image to PDF"],
        
        # Image Converter
        "input_img_label": "Select Input Image File(s):",
        "input_placeholder_single_img": "Select a single image file...",
        "input_placeholder_multi_img": "Select one or more image files...",
        "output_placeholder_img": "Select a destination folder...",
        "img_options_label": "Image Conversion Options:",
        "resolution_label": "Resolution:",
        "quality_label": "Quality:",
        "img_formats": ["JPEG", "PNG", "WEBP", "BMP", "GIF"],
        "img_qualities": ["Maximum (100)", "Very Good (95)", "Good (85)", "Medium (75)", "Low (50)"],
        
        # Video Converter
        "input_vid_label": "Select Input Video File(s):",
        "input_placeholder_single_vid": "Select a single video file...",
        "input_placeholder_multi_vid": "Select multiple video files for batch conversion...",
        "output_placeholder_vid": "Select a destination folder...",
        "vid_options_label": "Video Conversion Options:",
        "vid_formats": ["mp4", "mkv", "avi", "mov", "webm", "gif"],
        "vid_qualities": ["High", "Medium", "Low"],
        "shutdown_confirm_title": "Confirm PC Shutdown",
        "shutdown_confirm_msg": "Conversion is complete. Are you sure you want to shut down the PC now?",

        # Audio Converter
        "input_audio_label": "Select Input Audio File(s):",
        "input_placeholder_single_audio": "Select a single audio file...",
        "input_placeholder_multi_audio": "Select multiple audio files for batch conversion...",
        "output_placeholder_audio": "Select a destination folder...",
        "audio_options_label": "Audio Conversion Options:",
        "bitrate_label": "Bitrate:",
        "audio_formats": ["mp3", "wav", "aac", "flac", "ogg", "wma", "m4a"],
        
        # Video-Audio Converter (Extract Audio)
        "input_va_label": "Select Input Video File(s):",
        "output_va_label": "Select Audio Output Folder:",
        "input_placeholder_single_va": "Select a single video file...",
        "input_placeholder_multi_va": "Select multiple videos for batch conversion...",
        "output_placeholder_va": "Select audio destination folder...",
        "va_options_label": "Audio Extraction Options:",
        
        # Messages & Notifications
        "invalid_input_title": "Invalid Input",
        "invalid_input_file_msg": "Please select a valid input file.",
        "invalid_output_folder_msg": "Please select a valid output folder.",
        "no_ico_size_msg": "Please select at least one resolution size for the ICO file.",
        "conversion_logic_error": "Conversion logic for mode '{mode}' not found.",
        "docx_conv_error_title": "DOCX Conversion Error",
        "docx_conv_error_msg": "DOCX to PDF conversion failed. Make sure Microsoft Word (Windows) or LibreOffice (Linux/Mac) is installed.",
        "ffmpeg_not_found_title": "FFmpeg Not Found",
        "ffmpeg_not_found_msg": "ffmpeg was not found in the application directory or in the system PATH. Please place 'ffmpeg.exe' (Windows) or 'ffmpeg' (Mac/Linux) next to this application.",
        "batch_no_files_msg": "Please select at least one file for batch mode.",
        "conversion_success_msg": "Success! The conversion is complete.",
        "batch_complete_title": "Batch Complete",
        "batch_complete_msg": "Batch finished! All {count} files have been successfully converted.",
        "error_title": "Error",
        "error_during_conversion": "An error occurred: {error}",
        "error_in_batch": "Error on file {index}: {error}\n\nBatch process stopped.",
        "conversion_stopped": "Conversion stopped by user.",

        # Conversion Status
        "converting_page": "Converting page {current} of {total}...",
        "pdf_conversion_success": "Success! {total} PDF pages have been converted.",
        "opening_png": "Opening PNG file...",
        "saving_ico": "Saving to ICO format with selected sizes...",
        "done": "Done!",
        "png_to_ico_success": "Success! The PNG file has been converted to ICO.",
        "opening_image": "Opening {filename}...",
        "processing_image": "Processing image...",
        "saving_image": "Image saved successfully.",
        "image_conversion_success": "Success! The image has been converted to {format}.",
        "video_conversion_success": "Success! The video has been converted to {format}.",
        "getting_video_info": "Getting video info...",
        "converting_progress": "Converting... {progress}%",
        "getting_audio_info": "Getting audio info...",
        "audio_conversion_success": "Success! The audio has been converted to {format}.",
        "preparing_conversion": "Preparing conversion for {filename}...",
        "converting_batch_file": "Converting file {current} of {total}: {filename}",
        
        # About Tab
        "about_content": """<b>Macan Converter Pro</b>
Version 3.0 (Redesign)

Developed by: Macan Angkasa
© 2025 Danx Exodus

Macan Converter Pro is a modern multi-conversion application
for files, images, audio, and video.
It supports batch mode, various popular formats, and features a
clean and easy-to-use interface."""
    }
}


# --- Kelas Worker (Tidak ada perubahan signifikan pada logika internal) ---
class FileConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)
    
    def __init__(self, input_path, output_path, mode, out_format, ico_sizes=None, lang_dict=None):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.mode = mode
        self.out_format = out_format
        self.ico_sizes = ico_sizes if ico_sizes else []
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def run(self):
        try:
            mode_map = {
                "PDF ke Gambar": self._convert_pdf_to_image, "PDF to Image": self._convert_pdf_to_image,
                "PNG ke ICO": self._convert_png_to_ico, "PNG to ICO": self._convert_png_to_ico,
                "PDF ke TXT": self._convert_pdf_to_txt, "PDF to TXT": self._convert_pdf_to_txt,
                "PDF ke DOCX": self._convert_pdf_to_docx, "PDF to DOCX": self._convert_pdf_to_docx,
                "PDF ke XLSX": self._convert_pdf_to_xlsx, "PDF to XLSX": self._convert_pdf_to_xlsx,
                "DOCX ke PDF": self._convert_docx_to_pdf, "DOCX to PDF": self._convert_docx_to_pdf,
                "Gambar ke PDF": self._convert_image_to_pdf, "Image to PDF": self._convert_image_to_pdf,
            }
            conversion_func = mode_map.get(self.mode)
            if conversion_func:
                conversion_func()
            else:
                self.conversion_error.emit(self.lang["conversion_logic_error"].format(mode=self.mode))
        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False

    def _convert_pdf_to_image(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        for i in range(total_pages):
            if not self.is_running: break
            image = pdf[i].render(scale=200/72).to_pil()
            output_filename = os.path.join(self.output_path, f"{base_name}_page_{i+1}.{self.out_format.lower()}")
            image.save(output_filename)
            progress = int(((i + 1) / total_pages) * 100)
            status_text = self.lang["converting_page"].format(current=i+1, total=total_pages)
            self.progress_updated.emit(progress, status_text)
        if self.is_running:
            self.conversion_finished.emit(self.lang["pdf_conversion_success"].format(total=total_pages))

    def _convert_png_to_ico(self):
        self.progress_updated.emit(20, self.lang["opening_png"])
        if not self.ico_sizes:
            self.conversion_error.emit(self.lang["no_ico_size_msg"])
            return
        img = Image.open(self.input_path)
        self.progress_updated.emit(60, self.lang["saving_ico"])
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.ico")
        img.save(output_filename, format='ICO', sizes=self.ico_sizes)
        self.progress_updated.emit(100, self.lang["done"])
        self.conversion_finished.emit(self.lang["png_to_ico_success"])
    
    def _convert_pdf_to_txt(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.txt")
        full_text = ""
        for i in range(total_pages):
            if not self.is_running: break
            text_page = pdf[i].get_textpage()
            full_text += text_page.get_text_range() + "\n\n"
            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))
        with open(output_filename, "w", encoding="utf-8") as f:
            f.write(full_text)
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])
    
    def _convert_pdf_to_docx(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.docx")
        word_doc = docx.Document()
        for i in range(total_pages):
            if not self.is_running: break
            text_page = pdf[i].get_textpage()
            word_doc.add_paragraph(text_page.get_text_range())
            if i < total_pages - 1:
                word_doc.add_page_break()
            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))
        word_doc.save(output_filename)
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

    def _convert_pdf_to_xlsx(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.xlsx")
        workbook = openpyxl.Workbook()
        if "Sheet" in workbook.sheetnames:
            workbook.remove(workbook["Sheet"])
        for i in range(total_pages):
            if not self.is_running: break
            sheet = workbook.create_sheet(title=f"Page_{i+1}")
            page_text = pdf[i].get_textpage().get_text_range()
            lines = page_text.splitlines()
            for r, line in enumerate(lines, 1):
                sheet.cell(row=r, column=1, value=line)
            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))
        workbook.save(output_filename)
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

    def _convert_docx_to_pdf(self):
        self.progress_updated.emit(10, "Mempersiapkan konversi DOCX...")
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.pdf")
        try:
            self.progress_updated.emit(50, "Mengonversi file...")
            convert(self.input_path, output_filename)
            self.progress_updated.emit(100, self.lang["done"])
            self.conversion_finished.emit(self.lang["conversion_success_msg"])
        except Exception as e:
            self.conversion_error.emit(f"{self.lang['docx_conv_error_msg']} Detail: {e}")

    def _convert_image_to_pdf(self):
        total_files = len(self.input_path)
        if total_files == 0:
            self.conversion_error.emit(self.lang["batch_no_files_msg"])
            return
        base_name = os.path.splitext(os.path.basename(self.input_path[0]))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}_converted.pdf")
        images_pil = []
        for i, img_path in enumerate(self.input_path):
            if not self.is_running: break
            try:
                images_pil.append(Image.open(img_path).convert("RGB"))
            except Exception as e:
                self.conversion_error.emit(f"Gagal membuka gambar {os.path.basename(img_path)}: {e}")
                return
            progress = int(((i + 1) / total_files) * 100)
            self.progress_updated.emit(progress, f"Memproses gambar {i+1} dari {total_files}...")
        if self.is_running and images_pil:
            images_pil[0].save(output_filename, "PDF", resolution=100.0, save_all=True, append_images=images_pil[1:])
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

class ImageConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, input_path, output_path, out_format, resolution, quality_str, lang_dict=None):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.out_format = out_format
        self.resolution = resolution
        self.quality_str = quality_str
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def stop(self):
        self.is_running = False
        
    def run(self):
        try:
            self.progress_updated.emit(0, self.lang["opening_image"].format(filename=os.path.basename(self.input_path)))
            if not self.is_running: return
            
            img = Image.open(self.input_path)
            if self.resolution != "Original Size":
                res_parts = self.resolution.split(' ')[0].split('x')
                img = img.resize((int(res_parts[0]), int(res_parts[1])), Image.Resampling.LANCZOS)

            self.progress_updated.emit(50, self.lang["processing_image"])
            if not self.is_running: return

            base_name = os.path.splitext(os.path.basename(self.input_path))[0]
            output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format.lower()}")
            
            save_options = {}
            if self.out_format.lower() in ['jpeg', 'jpg', 'webp']:
                quality_map = {"Maximum (100)": 100, "Very Good (95)": 95, "Good (85)": 85, "Medium (75)": 75, "Low (50)": 50,
                               "Maksimum (100)": 100, "Sangat Baik (95)": 95, "Baik (85)": 85, "Sedang (75)": 75, "Rendah (50)": 50}
                save_options['quality'] = quality_map.get(self.quality_str, 85)
            
            if self.out_format.lower() in ['jpeg', 'jpg'] and img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')

            img.save(output_filename, **save_options)
            self.progress_updated.emit(100, self.lang["saving_image"])
            self.conversion_finished.emit(self.lang["image_conversion_success"].format(format=self.out_format.upper()))
        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

class VideoConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, ffmpeg_path, input_path, output_path, out_format, resolution, quality, lang_dict=None):
        super().__init__()
        self.ffmpeg_path = ffmpeg_path; self.input_path = input_path; self.output_path = output_path
        self.out_format = out_format; self.resolution = resolution; self.quality = quality
        self.is_running = True; self.lang = lang_dict if lang_dict else LANGUAGES["id"]
        self.process = None

    def _get_media_duration(self):
        command = [self.ffmpeg_path, '-i', self.input_path]
        try:
            startupinfo = subprocess.STARTUPINFO() if os.name == 'nt' else None
            if startupinfo: startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            result = subprocess.run(command, capture_output=True, text=True, stderr=subprocess.STDOUT, startupinfo=startupinfo)
            duration_search = re.search(r"Duration: (\d{2}):(\d{2}):(\d{2})\.(\d{2})", result.stdout)
            if duration_search:
                h, m, s = map(int, duration_search.groups()[:3])
                return (h * 3600) + (m * 60) + s
        except Exception: return None
        return None

    def run(self):
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format}")
        command = [self.ffmpeg_path, '-i', self.input_path]
        
        quality_map = {"Tinggi": "18", "Sedang": "23", "Rendah": "28", "High": "18", "Medium": "23", "Low": "28"}
        command.extend(['-crf', quality_map.get(self.quality, "23")])

        resolution_map = {"360p": "360", "480p": "480", "720p": "720", "1080p": "1080", "2K": "1440", "4K": "2160"}
        if self.resolution in resolution_map:
            command.extend(['-vf', f'scale=-2:{resolution_map[self.resolution]}'])

        command.extend(['-y', output_filename])
        
        try:
            self.progress_updated.emit(0, self.lang["getting_video_info"])
            total_duration = self._get_media_duration()
            
            startupinfo = subprocess.STARTUPINFO() if os.name == 'nt' else None
            if startupinfo: startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            self.process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, universal_newlines=True, encoding='utf-8', startupinfo=startupinfo)

            for line in iter(self.process.stdout.readline, ""):
                if not self.is_running: break
                if total_duration:
                    time_search = re.search(r"time=(\d{2}):(\d{2}):(\d{2})\.(\d{2})", line)
                    if time_search:
                        h, m, s = map(int, time_search.groups()[:3])
                        progress = int(((h*3600 + m*60 + s) / total_duration) * 100)
                        self.progress_updated.emit(progress, self.lang["converting_progress"].format(progress=progress))
            self.process.wait()

            if not self.is_running:
                self.conversion_error.emit(self.lang["conversion_stopped"])
            elif self.process.returncode == 0:
                self.conversion_finished.emit(self.lang["video_conversion_success"].format(format=self.out_format.upper()))
            else:
                self.conversion_error.emit(f"Failed. Exit code: {self.process.returncode}")
        except Exception as e:
            if self.is_running: self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False
        if self.process: self.process.terminate()

class AudioConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, ffmpeg_path, input_path, output_path, out_format, bitrate, lang_dict=None):
        super().__init__()
        self.ffmpeg_path = ffmpeg_path; self.input_path = input_path; self.output_path = output_path
        self.out_format = out_format; self.bitrate = bitrate
        self.is_running = True; self.lang = lang_dict if lang_dict else LANGUAGES["id"]
        self.process = None
        
    def _get_media_duration(self): # Sama seperti Video Worker
        command = [self.ffmpeg_path, '-i', self.input_path]
        try:
            startupinfo = subprocess.STARTUPINFO() if os.name == 'nt' else None
            if startupinfo: startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            result = subprocess.run(command, capture_output=True, text=True, stderr=subprocess.STDOUT, startupinfo=startupinfo)
            duration_search = re.search(r"Duration: (\d{2}):(\d{2}):(\d{2})\.(\d{2})", result.stdout)
            if duration_search:
                h, m, s = map(int, duration_search.groups()[:3])
                return (h * 3600) + (m * 60) + s
        except Exception: return None
        return None

    def run(self):
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format}")
        command = [self.ffmpeg_path, '-i', self.input_path, '-vn', '-b:a', self.bitrate.split(' ')[0] + 'k', '-y', output_filename]
        
        try:
            self.progress_updated.emit(0, self.lang["getting_audio_info"])
            total_duration = self._get_media_duration()
            
            startupinfo = subprocess.STARTUPINFO() if os.name == 'nt' else None
            if startupinfo: startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            self.process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, universal_newlines=True, encoding='utf-8', startupinfo=startupinfo)

            for line in iter(self.process.stdout.readline, ""):
                if not self.is_running: break
                if total_duration:
                    time_search = re.search(r"time=(\d{2}):(\d{2}):(\d{2})\.(\d{2})", line)
                    if time_search:
                        h, m, s = map(int, time_search.groups()[:3])
                        progress = int(((h*3600 + m*60 + s) / total_duration) * 100)
                        self.progress_updated.emit(progress, self.lang["converting_progress"].format(progress=progress))
            self.process.wait()

            if not self.is_running:
                self.conversion_error.emit(self.lang["conversion_stopped"])
            elif self.process.returncode == 0:
                self.conversion_finished.emit(self.lang["audio_conversion_success"].format(format=self.out_format.upper()))
            else:
                self.conversion_error.emit(f"Failed. Exit code: {self.process.returncode}")
        except Exception as e:
            if self.is_running: self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False
        if self.process: self.process.terminate()

# --- [REFACTOR] Kelas Utama Aplikasi dengan UI Baru ---
class RedesignedConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_lang = "id"
        self.lang = LANGUAGES[self.current_lang]
        
        self.setWindowTitle(self.lang["window_title"])
        self.setGeometry(100, 100, 800, 600)
        self.setMinimumSize(750, 550)
        
        # Setup path ikon aplikasi
        icon_path = "icon.ico"
        if hasattr(sys, "_MEIPASS"):
            icon_path = os.path.join(sys._MEIPASS, icon_path)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        self.thread = None
        self.worker = None
        
        # Variabel batch untuk setiap mode
        self.batch_files = {
            'document': [], 'image': [], 'video': [], 'audio': [], 'extract_audio': []
        }
        self.current_batch_index = 0
        
        self._setup_icons() # [BARU] Panggil setup ikon sebelum UI
        self._setup_ui()
        self._apply_stylesheet()

    # [PERBAIKAN] Metode _get_std_icon dihapus karena tidak andal.

    # [BARU] Metode untuk menyimpan data ikon SVG
    def _setup_icons(self):
        self.icons = {
            'video': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polygon points="5 3 19 12 5 21 5 3"></polygon></svg>""",
            'audio': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"></polygon><path d="M19.07 4.93a10 10 0 0 1 0 14.14M15.54 8.46a5 5 0 0 1 0 7.07"></path></svg>""",
            'extract_audio': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 19c-3.31 0-6-2.69-6-6V7c0-3.31 2.69-6 6-6s6 2.69 6 6v6c0 3.31-2.69 6-6 6z"></path><line x1="8" y1="22" x2="16" y2="22"></line><line x1="12" y1="1" x2="12" y2="4"></line></svg>""",
            'image': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"></rect><circle cx="8.5" cy="8.5" r="1.5"></circle><polyline points="21 15 16 10 5 21"></polyline></svg>""",
            'document': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"></path><polyline points="14 2 14 8 20 8"></polyline><line x1="16" y1="13" x2="8" y2="13"></line><line x1="16" y1="17" x2="8" y2="17"></line><polyline points="10 9 9 9 8 9"></polyline></svg>""",
            'about': """<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#ECEFF4" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"></circle><line x1="12" y1="16" x2="12" y2="12"></line><line x1="12" y1="8" x2="12.01" y2="8"></line></svg>"""
        }

    # [BARU] Helper function untuk membuat QIcon dari data SVG XML
    def _create_icon_from_svg(self, svg_data):
        byte_array = QByteArray(svg_data.encode('utf-8'))
        pixmap = QPixmap()
        pixmap.loadFromData(byte_array, 'SVG')
        return QIcon(pixmap)
        
    def _setup_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0,0,0,0)
        main_layout.setSpacing(0)

        # Header
        header_frame = QFrame()
        header_frame.setObjectName("headerFrame")
        header_layout = QHBoxLayout(header_frame)
        self.title_label = QLabel(self.lang["main_title"])
        self.title_label.setObjectName("titleLabel")
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["Indonesia", "English"])
        self.lang_combo.currentIndexChanged.connect(self.change_language)
        header_layout.addWidget(self.title_label)
        header_layout.addStretch()
        header_layout.addWidget(QLabel("Bahasa/Language:"))
        header_layout.addWidget(self.lang_combo)
        main_layout.addWidget(header_frame)
        
        # Toolbar
        toolbar_frame = QFrame()
        toolbar_frame.setObjectName("toolbarFrame")
        toolbar_layout = QHBoxLayout(toolbar_frame)
        self.output_folder_btn = QPushButton(self.lang["output_folder_btn"])
        self.output_folder_btn.clicked.connect(self.browse_master_output_folder)
        self.stop_btn = QPushButton(self.lang["stop_conversion_btn"])
        self.stop_btn.setObjectName("stopButton")
        self.stop_btn.clicked.connect(self.stop_conversion)
        self.stop_btn.setEnabled(False)
        self.start_btn = QPushButton(self.lang["start_conversion_btn"])
        self.start_btn.setObjectName("startButton")
        self.start_btn.clicked.connect(self.start_master_conversion)
        toolbar_layout.addWidget(self.output_folder_btn)
        toolbar_layout.addStretch()
        toolbar_layout.addWidget(self.stop_btn)
        toolbar_layout.addWidget(self.start_btn)
        main_layout.addWidget(toolbar_frame)
        
        # Content (Sidebar + StackedWidget)
        content_layout = QHBoxLayout()
        content_layout.setSpacing(0)
        
        # Sidebar/Category List
        self.category_list = QListWidget()
        self.category_list.setObjectName("categoryList")
        self.category_list.setFixedWidth(180)
        self.category_list.setIconSize(QSize(32, 32))
        self.category_list.currentRowChanged.connect(self.change_category)
        content_layout.addWidget(self.category_list)
        
        # Content Stack
        self.stacked_widget = QStackedWidget()
        self.stacked_widget.setObjectName("stackedWidget")
        
        # Buat halaman untuk setiap kategori
        self.video_page = self._create_video_page()
        self.audio_page = self._create_audio_page()
        self.extract_audio_page = self._create_extract_audio_page()
        self.image_page = self._create_image_page()
        self.document_page = self._create_document_page()
        self.about_page = self._create_about_page()

        self.stacked_widget.addWidget(self.video_page)
        self.stacked_widget.addWidget(self.audio_page)
        self.stacked_widget.addWidget(self.extract_audio_page)
        self.stacked_widget.addWidget(self.image_page)
        self.stacked_widget.addWidget(self.document_page)
        self.stacked_widget.addWidget(self.about_page)

        self.populate_categories() # Panggil setelah page dibuat
        
        content_layout.addWidget(self.stacked_widget)
        main_layout.addLayout(content_layout, 1)

        # Status Bar
        status_bar_frame = QFrame()
        status_bar_frame.setObjectName("statusBarFrame")
        status_bar_layout = QHBoxLayout(status_bar_frame)
        self.status_label = QLabel(self.lang["ready_status"])
        self.status_label.setObjectName("statusLabel")
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)
        self.shutdown_checkbox = QCheckBox(self.lang["shutdown_checkbox"])
        status_bar_layout.addWidget(self.status_label, 1)
        status_bar_layout.addWidget(self.progress_bar, 1)
        status_bar_layout.addSpacing(20)
        status_bar_layout.addWidget(self.shutdown_checkbox)
        
        main_layout.addWidget(status_bar_frame)

    def populate_categories(self):
        self.category_list.clear()
        # [PERBAIKAN] Menggunakan sistem ikon SVG yang baru dan andal
        # Urutan harus sama dengan urutan penambahan widget ke stacked_widget
        categories = [
            (self.lang["tab_video"], self.icons['video']),
            (self.lang["tab_audio"], self.icons['audio']),
            (self.lang["tab_video_audio"], self.icons['extract_audio']),
            (self.lang["tab_image"], self.icons['image']),
            (self.lang["tab_file"], self.icons['document']),
            (self.lang["tab_about"], self.icons['about'])
        ]
        
        for name, svg_data in categories:
            icon = self._create_icon_from_svg(svg_data)
            item = QListWidgetItem(icon, name)
            item.setSizeHint(QSize(self.category_list.width(), 40))
            self.category_list.addItem(item)
            
        self.category_list.setCurrentRow(0)


    def change_category(self, index):
        self.stacked_widget.setCurrentIndex(index)

    def _create_video_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(15)

        # Input
        self.vid_input_label = QLabel(self.lang["input_vid_label"])
        layout.addWidget(self.vid_input_label)
        input_layout = QHBoxLayout()
        self.vid_input_path_edit = QLineEdit(); self.vid_input_path_edit.setReadOnly(True)
        self.vid_browse_btn = QPushButton(self.lang["browse_btn"])
        self.vid_browse_btn.clicked.connect(self.browse_video_input_file)
        input_layout.addWidget(self.vid_input_path_edit); input_layout.addWidget(self.vid_browse_btn)
        layout.addLayout(input_layout)
        self.vid_output_path_edit = QLineEdit(); self.vid_output_path_edit.setPlaceholderText(self.lang["output_placeholder_vid"])
        self.vid_output_path_edit.setReadOnly(True) # Diisi oleh tombol utama
        layout.addWidget(self.vid_output_path_edit)
        self.batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.batch_mode_checkbox.stateChanged.connect(self._update_video_input_ui)
        layout.addWidget(self.batch_mode_checkbox)
        
        # Options
        self.vid_options_label = QLabel(f"<b>{self.lang['vid_options_label']}</b>")
        layout.addWidget(self.vid_options_label)
        options_layout = QHBoxLayout()
        self.vid_format_label = QLabel(self.lang["output_format_label"])
        self.vid_format_combo = QComboBox(); self.vid_format_combo.addItems(self.lang["vid_formats"])
        self.vid_res_label = QLabel(self.lang["resolution_label"])
        self.vid_resolution_combo = QComboBox(); self.vid_resolution_combo.addItems(["Original Size", "360p", "480p", "720p", "1080p", "2K", "4K"])
        self.vid_quality_label = QLabel(self.lang["quality_label"])
        self.vid_quality_combo = QComboBox(); self.vid_quality_combo.addItems(self.lang["vid_qualities"])
        self.vid_quality_combo.setCurrentText(self.lang["vid_qualities"][1])
        options_layout.addWidget(self.vid_format_label); options_layout.addWidget(self.vid_format_combo, 1)
        options_layout.addWidget(self.vid_res_label); options_layout.addWidget(self.vid_resolution_combo, 1)
        options_layout.addWidget(self.vid_quality_label); options_layout.addWidget(self.vid_quality_combo, 1)
        layout.addLayout(options_layout)
        
        layout.addStretch()
        self._update_video_input_ui()
        return page

    def _create_audio_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(15)

        self.audio_input_label = QLabel(self.lang["input_audio_label"])
        layout.addWidget(self.audio_input_label)
        input_layout = QHBoxLayout()
        self.audio_input_path_edit = QLineEdit(); self.audio_input_path_edit.setReadOnly(True)
        self.audio_browse_btn = QPushButton(self.lang["browse_btn"])
        self.audio_browse_btn.clicked.connect(self.browse_audio_input_file)
        input_layout.addWidget(self.audio_input_path_edit); input_layout.addWidget(self.audio_browse_btn)
        layout.addLayout(input_layout)
        self.audio_output_path_edit = QLineEdit(); self.audio_output_path_edit.setPlaceholderText(self.lang["output_placeholder_audio"]); self.audio_output_path_edit.setReadOnly(True)
        layout.addWidget(self.audio_output_path_edit)
        self.audio_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.audio_batch_mode_checkbox.stateChanged.connect(self._update_audio_input_ui)
        layout.addWidget(self.audio_batch_mode_checkbox)

        self.audio_options_label = QLabel(f"<b>{self.lang['audio_options_label']}</b>")
        layout.addWidget(self.audio_options_label)
        options_layout = QHBoxLayout()
        self.audio_format_label = QLabel(self.lang["output_format_label"])
        self.audio_format_combo = QComboBox(); self.audio_format_combo.addItems(self.lang["audio_formats"])
        self.audio_bitrate_label = QLabel(self.lang["bitrate_label"])
        self.audio_bitrate_combo = QComboBox(); self.audio_bitrate_combo.addItems(["96 kbps", "128 kbps", "192 kbps", "256 kbps", "320 kbps"]); self.audio_bitrate_combo.setCurrentText("192 kbps")
        options_layout.addWidget(self.audio_format_label); options_layout.addWidget(self.audio_format_combo, 1)
        options_layout.addWidget(self.audio_bitrate_label); options_layout.addWidget(self.audio_bitrate_combo, 1)
        options_layout.addStretch()
        layout.addLayout(options_layout)

        layout.addStretch()
        self._update_audio_input_ui()
        return page

    def _create_extract_audio_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(15)
        
        self.va_input_label = QLabel(self.lang["input_va_label"])
        layout.addWidget(self.va_input_label)
        input_layout = QHBoxLayout()
        self.va_input_path_edit = QLineEdit(); self.va_input_path_edit.setReadOnly(True)
        self.va_browse_btn = QPushButton(self.lang["browse_btn"]); self.va_browse_btn.clicked.connect(self.browse_va_input_file)
        input_layout.addWidget(self.va_input_path_edit); input_layout.addWidget(self.va_browse_btn)
        layout.addLayout(input_layout)
        self.va_output_path_edit = QLineEdit(); self.va_output_path_edit.setPlaceholderText(self.lang["output_placeholder_va"]); self.va_output_path_edit.setReadOnly(True)
        layout.addWidget(self.va_output_path_edit)
        self.va_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"]); self.va_batch_mode_checkbox.stateChanged.connect(self._update_va_input_ui)
        layout.addWidget(self.va_batch_mode_checkbox)

        self.va_options_label = QLabel(f"<b>{self.lang['va_options_label']}</b>")
        layout.addWidget(self.va_options_label)
        options_layout = QHBoxLayout()
        self.va_format_label = QLabel(self.lang["output_format_label"])
        self.va_format_combo = QComboBox(); self.va_format_combo.addItems(self.lang["audio_formats"])
        self.va_bitrate_label = QLabel(self.lang["bitrate_label"])
        self.va_bitrate_combo = QComboBox(); self.va_bitrate_combo.addItems(["96 kbps", "128 kbps", "192 kbps", "256 kbps", "320 kbps"]); self.va_bitrate_combo.setCurrentText("192 kbps")
        options_layout.addWidget(self.va_format_label); options_layout.addWidget(self.va_format_combo, 1)
        options_layout.addWidget(self.va_bitrate_label); options_layout.addWidget(self.va_bitrate_combo, 1)
        options_layout.addStretch()
        layout.addLayout(options_layout)
        
        layout.addStretch()
        self._update_va_input_ui()
        return page

    def _create_image_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(15)
        
        self.img_input_label = QLabel(self.lang["input_img_label"])
        layout.addWidget(self.img_input_label)
        input_layout = QHBoxLayout()
        self.img_input_path_edit = QLineEdit(); self.img_input_path_edit.setReadOnly(True)
        self.img_browse_btn = QPushButton(self.lang["browse_btn"]); self.img_browse_btn.clicked.connect(self.browse_image_input_file)
        input_layout.addWidget(self.img_input_path_edit); input_layout.addWidget(self.img_browse_btn)
        layout.addLayout(input_layout)
        self.img_output_path_edit = QLineEdit(); self.img_output_path_edit.setPlaceholderText(self.lang["output_placeholder_img"]); self.img_output_path_edit.setReadOnly(True)
        layout.addWidget(self.img_output_path_edit)
        self.img_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"]); self.img_batch_mode_checkbox.stateChanged.connect(self._update_image_input_ui)
        layout.addWidget(self.img_batch_mode_checkbox)

        self.img_options_label = QLabel(f"<b>{self.lang['img_options_label']}</b>")
        layout.addWidget(self.img_options_label)
        options_layout = QHBoxLayout()
        self.img_format_label = QLabel(self.lang["output_format_label"])
        self.img_format_combo = QComboBox(); self.img_format_combo.addItems(self.lang["img_formats"]); self.img_format_combo.currentTextChanged.connect(self._update_image_options)
        self.img_res_label = QLabel(self.lang["resolution_label"])
        self.img_resolution_combo = QComboBox(); self.img_resolution_combo.addItems(["Original Size", "320x240", "640x480", "800x600", "1280x720 (HD)", "1920x1080 (Full HD)", "2560x1440 (2K)", "3840x2160 (4K)"])
        self.img_quality_label = QLabel(self.lang["quality_label"])
        self.img_quality_combo = QComboBox(); self.img_quality_combo.addItems(self.lang["img_qualities"]); self.img_quality_combo.setCurrentText(self.lang["img_qualities"][2])
        options_layout.addWidget(self.img_format_label); options_layout.addWidget(self.img_format_combo, 1)
        options_layout.addWidget(self.img_res_label); options_layout.addWidget(self.img_resolution_combo, 1)
        options_layout.addWidget(self.img_quality_label); options_layout.addWidget(self.img_quality_combo, 1)
        layout.addLayout(options_layout)

        layout.addStretch()
        self._update_image_input_ui(); self._update_image_options()
        return page

    def _create_document_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20); layout.setSpacing(15)
        
        self.file_input_label = QLabel(self.lang["input_label"])
        layout.addWidget(self.file_input_label)
        input_layout = QHBoxLayout()
        self.input_path_edit = QLineEdit(); self.input_path_edit.setReadOnly(True)
        self.browse_input_btn = QPushButton(self.lang["browse_btn"]); self.browse_input_btn.clicked.connect(self.browse_input_file)
        input_layout.addWidget(self.input_path_edit); input_layout.addWidget(self.browse_input_btn)
        layout.addLayout(input_layout)
        self.output_path_edit = QLineEdit(); self.output_path_edit.setPlaceholderText(self.lang["output_placeholder_folder"]); self.output_path_edit.setReadOnly(True)
        layout.addWidget(self.output_path_edit)
        
        self.conv_mode_label = QLabel(f"<b>{self.lang['conversion_mode_label']}</b>")
        layout.addWidget(self.conv_mode_label)
        settings_layout = QHBoxLayout()
        self.mode_label = QLabel(self.lang["mode_label"])
        self.mode_combo = QComboBox(); self.mode_combo.addItems(self.lang["modes"]); self.mode_combo.currentTextChanged.connect(self._update_ui_for_mode)
        self.format_label = QLabel(self.lang["output_format_label"])
        self.format_combo = QComboBox(); self.format_combo.addItems(["PNG", "JPG", "WEBP"])
        settings_layout.addWidget(self.mode_label); settings_layout.addWidget(self.mode_combo, 1)
        settings_layout.addWidget(self.format_label); settings_layout.addWidget(self.format_combo, 1)
        layout.addLayout(settings_layout)
        
        self.ico_sizes_frame = QFrame(); self.ico_sizes_frame.setObjectName("icoFrame")
        ico_layout = QVBoxLayout(self.ico_sizes_frame)
        self.ico_res_label = QLabel(self.lang["ico_res_label"])
        ico_layout.addWidget(self.ico_res_label)
        ico_grid = QGridLayout()
        self.ico_checkboxes = {}
        sizes = ["16x16", "32x32", "48x48", "64x64", "128x128", "256x256"]
        for i, size in enumerate(sizes):
            checkbox = QCheckBox(size); checkbox.setChecked(True)
            self.ico_checkboxes[size] = checkbox
            ico_grid.addWidget(checkbox, i // 3, i % 3)
        ico_layout.addLayout(ico_grid)
        layout.addWidget(self.ico_sizes_frame)

        layout.addStretch()
        self._update_ui_for_mode(self.mode_combo.currentText())
        return page

    def _create_about_page(self):
        page = QWidget()
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        self.about_text = QTextEdit()
        self.about_text.setReadOnly(True)
        self.about_text.setHtml(self.lang["about_content"])
        self.about_text.setObjectName("aboutText")
        layout.addWidget(self.about_text)
        return page

    def _apply_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow, QWidget {
                background-color: #252525; /* Abu-abu sangat gelap untuk background utama */
                color: #E0E0E0;
                font-family: Segoe UI, sans-serif;
            }

            /* [PERUBAHAN] Menghilangkan background header dan status bar */
            #headerFrame, #statusBarFrame {
                background-color: transparent;
                border: none; /* Menghapus border atas/bawah */
            }

            /* Toolbar tetap memiliki background terpisah */
            #toolbarFrame {
                background-color: #333333; /* Abu-abu sedikit lebih terang */
                border-bottom: 1px solid #444444;
            }
            
            #headerFrame { padding: 5px 10px; }
            #toolbarFrame { padding: 8px 10px; }
            #statusBarFrame { padding: 5px 10px; }
            
            #titleLabel { font-size: 14pt; font-weight: bold; color: #FFFFFF; }
            #statusLabel { color: #A3BE8C; font-size: 9pt; }
            QLabel { font-size: 10pt; }
            
            #categoryList {
                background-color: #333333; 
                border: none; 
                border-right: 1px solid #444444;
            }
            #categoryList::item {
                padding: 10px; 
                border-bottom: 1px solid #444444; 
                color: #E0E0E0;
            }
            #categoryList::item:selected { 
                background-color: #5A5A5A;
                color: #FFFFFF; 
            }
            #categoryList::item:hover:!selected { 
                background-color: #4A4A4A; 
            }
            
            #stackedWidget > QWidget { 
                background-color: #252525; 
            }
            
            QLineEdit, QComboBox, QTextEdit {
                background-color: #3A3A3A; 
                border: 1px solid #555555; 
                border-radius: 4px;
                padding: 6px; 
                color: #E0E0E0;
            }
            QPushButton {
                background-color: #555555; 
                color: #FFFFFF; 
                border: 1px solid #666666;
                padding: 8px 16px; 
                border-radius: 4px; 
                font-size: 10pt;
            }
            QPushButton:hover { background-color: #6A6A6A; }
            QPushButton:pressed { background-color: #444444; }
            
            #startButton { background-color: #A3BE8C; border: 1px solid #A3BE8C; color: #2E3440; font-weight: bold; }
            #startButton:hover { background-color: #b7d1a4; }
            #stopButton { background-color: #BF616A; border: 1px solid #BF616A; color: #ECEFF4; font-weight: bold; }
            #stopButton:hover { background-color: #d08770; }
            #startButton:disabled, #stopButton:disabled { background-color: #404040; color: #6F6F6F; border-color: #404040; }
            
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView { background-color: #3A3A3A; border: 1px solid #666666; color: #E0E0E0; }

            QProgressBar {
                border: none; 
                background-color: #3A3A3A; 
                border-radius: 4px; 
                text-align: center; 
                height: 12px;
            }
            QProgressBar::chunk { 
                background-color: #A3BE8C;
                border-radius: 4px; 
            }
            
            #icoFrame { border: 1px solid #444444; border-radius: 5px; padding: 10px; margin-top: 10px;}
            #aboutText { background-color: #333333; border: 1px solid #444444; color: #E0E0E0; font-size: 11pt; padding: 10px; }
            
            QCheckBox { font-size: 9pt; padding-top: 5px; }
            QCheckBox::indicator { border: 1px solid #777777; background-color: #444444; border-radius: 3px; width: 14px; height: 14px; }
            QCheckBox::indicator:hover { border: 1px solid #999999; }
            QCheckBox::indicator:checked { background-color: #A3BE8C; border: 1px solid #A3BE8C; }
        """)

    def change_language(self):
        self.current_lang = "en" if self.lang_combo.currentText() == "English" else "id"
        self.lang = LANGUAGES[self.current_lang]
        self._retranslate_ui()

    def _retranslate_ui(self):
        self.setWindowTitle(self.lang["window_title"])
        self.title_label.setText(self.lang["main_title"])
        self.output_folder_btn.setText(self.lang["output_folder_btn"])
        self.start_btn.setText(self.lang["start_conversion_btn"])
        self.stop_btn.setText(self.lang["stop_conversion_btn"])
        self.status_label.setText(self.lang["ready_status"])
        self.shutdown_checkbox.setText(self.lang["shutdown_checkbox"])
        self.populate_categories()

        # Video
        self.vid_input_label.setText(self.lang["input_vid_label"])
        self._update_video_input_ui()
        self.vid_browse_btn.setText(self.lang["browse_btn"])
        self.vid_output_path_edit.setPlaceholderText(self.lang["output_placeholder_vid"])
        self.batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.vid_options_label.setText(f"<b>{self.lang['vid_options_label']}</b>")
        self.vid_format_label.setText(self.lang["output_format_label"])
        self.vid_format_combo.clear(); self.vid_format_combo.addItems(self.lang["vid_formats"])
        self.vid_res_label.setText(self.lang["resolution_label"])
        self.vid_quality_label.setText(self.lang["quality_label"])
        self.vid_quality_combo.clear(); self.vid_quality_combo.addItems(self.lang["vid_qualities"])
        self.vid_quality_combo.setCurrentText(self.lang["vid_qualities"][1])
        
        # Audio
        self.audio_input_label.setText(self.lang["input_audio_label"])
        self._update_audio_input_ui()
        self.audio_browse_btn.setText(self.lang["browse_btn"])
        self.audio_output_path_edit.setPlaceholderText(self.lang["output_placeholder_audio"])
        self.audio_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.audio_options_label.setText(f"<b>{self.lang['audio_options_label']}</b>")
        self.audio_format_label.setText(self.lang["output_format_label"])
        self.audio_format_combo.clear(); self.audio_format_combo.addItems(self.lang["audio_formats"])
        self.audio_bitrate_label.setText(self.lang["bitrate_label"])

        # Extract Audio
        self.va_input_label.setText(self.lang["input_va_label"])
        self._update_va_input_ui()
        self.va_browse_btn.setText(self.lang["browse_btn"])
        self.va_output_path_edit.setPlaceholderText(self.lang["output_placeholder_va"])
        self.va_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.va_options_label.setText(f"<b>{self.lang['va_options_label']}</b>")
        self.va_format_label.setText(self.lang["output_format_label"])
        self.va_format_combo.clear(); self.va_format_combo.addItems(self.lang["audio_formats"])
        self.va_bitrate_label.setText(self.lang["bitrate_label"])

        # Image
        self.img_input_label.setText(self.lang["input_img_label"])
        self._update_image_input_ui()
        self.img_browse_btn.setText(self.lang["browse_btn"])
        self.img_output_path_edit.setPlaceholderText(self.lang["output_placeholder_img"])
        self.img_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.img_options_label.setText(f"<b>{self.lang['img_options_label']}</b>")
        self.img_format_label.setText(self.lang["output_format_label"])
        self.img_format_combo.clear(); self.img_format_combo.addItems(self.lang["img_formats"])
        self.img_res_label.setText(self.lang["resolution_label"])
        self.img_quality_label.setText(self.lang["quality_label"])
        self.img_quality_combo.clear(); self.img_quality_combo.addItems(self.lang["img_qualities"])
        self.img_quality_combo.setCurrentText(self.lang["img_qualities"][2])

        # Document
        self.file_input_label.setText(self.lang["input_label"])
        self.input_path_edit.setPlaceholderText(self.lang["input_placeholder_file"])
        self.browse_input_btn.setText(self.lang["browse_btn"])
        self.output_path_edit.setPlaceholderText(self.lang["output_placeholder_folder"])
        self.conv_mode_label.setText(f"<b>{self.lang['conversion_mode_label']}</b>")
        self.mode_label.setText(self.lang["mode_label"])
        self.mode_combo.clear(); self.mode_combo.addItems(self.lang["modes"])
        self.format_label.setText(self.lang["output_format_label"])
        self.ico_res_label.setText(self.lang["ico_res_label"])

        # About
        self.about_text.setHtml(self.lang["about_content"])
        
    def _find_ffmpeg(self):
        base_path = sys._MEIPASS if hasattr(sys, "_MEIPASS") else os.path.dirname(os.path.abspath(__file__))
        local_ffmpeg = os.path.join(base_path, "ffmpeg.exe" if os.name == 'nt' else "ffmpeg")
        if os.path.exists(local_ffmpeg): return local_ffmpeg
        from shutil import which
        return "ffmpeg" if which("ffmpeg") else None

    # --- Master Start/Stop/Browse Logic ---
    def start_master_conversion(self):
        # Panggil fungsi start yang sesuai dengan halaman aktif
        current_page_index = self.stacked_widget.currentIndex()
        page_map = {
            0: self.start_video_conversion,
            1: self.start_audio_conversion,
            2: self.start_va_conversion,
            3: self.start_image_conversion,
            4: self.start_file_conversion,
        }
        start_function = page_map.get(current_page_index)
        if start_function:
            start_function()

    def stop_conversion(self):
        if self.worker:
            self.worker.stop()
        self.stop_btn.setEnabled(False)
        self.start_btn.setEnabled(True)

    def browse_master_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_label"])
        if not folder_path: return

        # Update QLineEdit di halaman yang sedang aktif
        current_widget = self.stacked_widget.currentWidget()
        output_edit = current_widget.findChild(QLineEdit, "output_path_edit") # Asumsi nama object konsisten
        
        # Cari berdasarkan nama variabel jika objectName tidak diset
        if not output_edit:
            edits = current_widget.findChildren(QLineEdit)
            # Dapatkan yang placeholder-nya berisi 'output'
            output_edits = [e for e in edits if 'output' in e.placeholderText().lower() or 'tujuan' in e.placeholderText().lower()]
            if output_edits:
                output_edit = output_edits[0]

        if output_edit:
            output_edit.setText(folder_path)
        else: # Fallback untuk semua halaman
            self.vid_output_path_edit.setText(folder_path)
            self.audio_output_path_edit.setText(folder_path)
            self.va_output_path_edit.setText(folder_path)
            self.img_output_path_edit.setText(folder_path)
            self.output_path_edit.setText(folder_path)

    # --- Logika untuk File Converter ---
    def browse_input_file(self):
        mode = self.mode_combo.currentText()
        filters = {"PDF ke Gambar": "PDF Files (*.pdf)", "PNG ke ICO": "PNG Files (*.png)", "PDF ke TXT": "PDF Files (*.pdf)",
                   "PDF ke DOCX": "PDF Files (*.pdf)", "PDF ke XLSX": "PDF Files (*.pdf)", "DOCX ke PDF": "Word Documents (*.docx)",
                   "Gambar ke PDF": "Image Files (*.png *.jpg *.jpeg *.bmp *.webp *.gif)",
                   # Terjemahan EN
                   "PDF to Image": "PDF Files (*.pdf)", "PNG to ICO": "PNG Files (*.png)", "PDF to TXT": "PDF Files (*.pdf)",
                   "PDF to DOCX": "PDF Files (*.pdf)", "PDF to XLSX": "PDF Files (*.pdf)", "DOCX to PDF": "Word Documents (*.docx)",
                   "Image to PDF": "Image Files (*.png *.jpg *.jpeg *.bmp *.webp *.gif)"}
        filter_str = filters.get(mode, "All Files (*)")
        is_multi_file_mode = mode in ["Gambar ke PDF", "Image to PDF"]

        if is_multi_file_mode:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_label"], "", filter_str)
            if files: self.batch_files['document'] = files; self.input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            self.batch_files['document'] = []
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_label"], "", filter_str)
            if file_path: self.input_path_edit.setText(file_path)

    def _update_ui_for_mode(self, mode):
        is_pdf_to_image = mode in ["PDF ke Gambar", "PDF to Image"]
        is_png_to_ico = mode in ["PNG ke ICO", "PNG to ICO"]
        is_image_to_pdf = mode in ["Gambar ke PDF", "Image to PDF"]
        self.format_combo.setVisible(is_pdf_to_image); self.format_label.setVisible(is_pdf_to_image)
        self.ico_sizes_frame.setVisible(is_png_to_ico)
        self.input_path_edit.clear(); self.batch_files['document'] = []
        placeholder = "Pilih satu atau lebih file gambar..." if is_image_to_pdf else self.lang["input_placeholder_file"]
        self.input_path_edit.setPlaceholderText(placeholder)
            
    def start_file_conversion(self):
        output_path = self.output_path_edit.text()
        mode = self.mode_combo.currentText()
        is_multi_file_mode = mode in ["Gambar ke PDF", "Image to PDF"]
        input_path = self.batch_files['document'] if is_multi_file_mode else self.input_path_edit.text()

        if not input_path or (not is_multi_file_mode and not os.path.exists(input_path)):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"]); return
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"]); return

        ico_sizes = []
        if mode in ["PNG ke ICO", "PNG to ICO"]:
            ico_sizes = [(int(s.split('x')[0]), int(s.split('x')[1])) for s, cb in self.ico_checkboxes.items() if cb.isChecked()]
            if not ico_sizes: QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["no_ico_size_msg"]); return

        self._set_ui_for_conversion(True)
        self.worker = FileConversionWorker(input_path, output_path, mode, self.format_combo.currentText(), ico_sizes, self.lang)
        self._run_worker()
        
    # --- Logika untuk Image Converter ---
    def _update_image_input_ui(self):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        placeholder = self.lang["input_placeholder_multi_img"] if is_batch else self.lang["input_placeholder_single_img"]
        self.img_input_path_edit.setPlaceholderText(placeholder)
        self.img_input_path_edit.clear(); self.batch_files['image'] = []

    def _update_image_options(self):
        has_quality = self.img_format_combo.currentText().lower() in ['jpeg', 'jpg', 'webp']
        self.img_quality_label.setVisible(has_quality); self.img_quality_combo.setVisible(has_quality)

    def browse_image_input_file(self):
        filter_str = "Image Files (*.png *.jpg *.jpeg *.bmp *.webp *.gif)"
        if self.img_batch_mode_checkbox.isChecked():
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_img_label"], "", filter_str)
            if files: self.batch_files['image'] = files; self.img_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_img_label"], "", filter_str)
            if file_path: self.img_input_path_edit.setText(file_path)

    def start_image_conversion(self):
        self._start_batch_or_single('image', self.img_output_path_edit, self.img_batch_mode_checkbox, self.batch_files['image'], self.img_input_path_edit, self._start_single_image_conversion)

    def _start_single_image_conversion(self, input_path):
        self._set_ui_for_conversion(True)
        self.status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))
        self.worker = ImageConversionWorker(input_path, self.img_output_path_edit.text(), self.img_format_combo.currentText(), self.img_resolution_combo.currentText(), self.img_quality_combo.currentText(), self.lang)
        self._run_worker()

    # --- Logika untuk Video Converter ---
    def _update_video_input_ui(self):
        is_batch = self.batch_mode_checkbox.isChecked()
        placeholder = self.lang["input_placeholder_multi_vid"] if is_batch else self.lang["input_placeholder_single_vid"]
        self.vid_input_path_edit.setPlaceholderText(placeholder)
        self.vid_input_path_edit.clear(); self.batch_files['video'] = []

    def browse_video_input_file(self):
        filter_str = "Video Files (*.mp4 *.mkv *.avi *.mov *.flv *.wmv)"
        if self.batch_mode_checkbox.isChecked():
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_vid_label"], "", filter_str)
            if files: self.batch_files['video'] = files; self.vid_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_vid_label"], "", filter_str)
            if file_path: self.vid_input_path_edit.setText(file_path)

    def start_video_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path: QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"]); return
        self._start_batch_or_single('video', self.vid_output_path_edit, self.batch_mode_checkbox, self.batch_files['video'], self.vid_input_path_edit, self._start_single_video_conversion)

    def _start_single_video_conversion(self, input_path):
        self._set_ui_for_conversion(True)
        self.status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))
        self.worker = VideoConversionWorker(self.ffmpeg_path, input_path, self.vid_output_path_edit.text(), self.vid_format_combo.currentText(), self.vid_resolution_combo.currentText(), self.vid_quality_combo.currentText(), self.lang)
        self._run_worker()

    # --- Logika untuk Audio Converter ---
    def _update_audio_input_ui(self):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        placeholder = self.lang["input_placeholder_multi_audio"] if is_batch else self.lang["input_placeholder_single_audio"]
        self.audio_input_path_edit.setPlaceholderText(placeholder)
        self.audio_input_path_edit.clear(); self.batch_files['audio'] = []

    def browse_audio_input_file(self):
        filter_str = "Audio Files (*.mp3 *.wav *.aac *.flac *.ogg *.wma *.m4a)"
        if self.audio_batch_mode_checkbox.isChecked():
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_audio_label"], "", filter_str)
            if files: self.batch_files['audio'] = files; self.audio_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_audio_label"], "", filter_str)
            if file_path: self.audio_input_path_edit.setText(file_path)

    def start_audio_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path: QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"]); return
        self._start_batch_or_single('audio', self.audio_output_path_edit, self.audio_batch_mode_checkbox, self.batch_files['audio'], self.audio_input_path_edit, self._start_single_audio_conversion)

    def _start_single_audio_conversion(self, input_path):
        self._set_ui_for_conversion(True)
        self.status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))
        self.worker = AudioConversionWorker(self.ffmpeg_path, input_path, self.audio_output_path_edit.text(), self.audio_format_combo.currentText(), self.audio_bitrate_combo.currentText(), self.lang)
        self._run_worker()
            
    # --- Logika untuk Video-Audio Converter ---
    def _update_va_input_ui(self):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        placeholder = self.lang["input_placeholder_multi_va"] if is_batch else self.lang["input_placeholder_single_va"]
        self.va_input_path_edit.setPlaceholderText(placeholder)
        self.va_input_path_edit.clear(); self.batch_files['extract_audio'] = []

    def browse_va_input_file(self):
        filter_str = "Video Files (*.mp4 *.mkv *.avi *.mov *.flv *.wmv)"
        if self.va_batch_mode_checkbox.isChecked():
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_va_label"], "", filter_str)
            if files: self.batch_files['extract_audio'] = files; self.va_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_va_label"], "", filter_str)
            if file_path: self.va_input_path_edit.setText(file_path)

    def start_va_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path: QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"]); return
        self._start_batch_or_single('extract_audio', self.va_output_path_edit, self.va_batch_mode_checkbox, self.batch_files['extract_audio'], self.va_input_path_edit, self._start_single_va_conversion)

    def _start_single_va_conversion(self, input_path):
        self._set_ui_for_conversion(True)
        self.status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))
        self.worker = AudioConversionWorker(self.ffmpeg_path, input_path, self.va_output_path_edit.text(), self.va_format_combo.currentText(), self.va_bitrate_combo.currentText(), self.lang)
        self._run_worker()
    
    # --- [REFACTOR] Helper Terpusat untuk Batch/Single ---
    def _start_batch_or_single(self, mode_key, output_edit, batch_checkbox, file_list, input_edit, single_start_func):
        output_path = output_edit.text()
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"]); return

        self.current_mode_key = mode_key
        self.current_single_start_func = single_start_func

        if batch_checkbox.isChecked():
            if not file_list:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["batch_no_files_msg"]); return
            self.current_batch_index = 0
            self._start_next_batch_conversion()
        else:
            input_path = input_edit.text()
            if not input_path or not os.path.exists(input_path):
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"]); return
            single_start_func(input_path)

    def _start_next_batch_conversion(self):
        file_list = self.batch_files[self.current_mode_key]
        if self.current_batch_index < len(file_list):
            input_path = file_list[self.current_batch_index]
            self.status_label.setText(
                self.lang["converting_batch_file"].format(
                    current=self.current_batch_index + 1, total=len(file_list), filename=os.path.basename(input_path)
                )
            )
            self.current_single_start_func(input_path)
        else: # Batch selesai
            final_msg = self.lang["batch_complete_msg"].format(count=len(file_list))
            self.status_label.setText(final_msg)
            self._set_ui_for_conversion(False)
            QMessageBox.information(self, self.lang["batch_complete_title"], final_msg)
            self._open_output_folder(self.get_active_output_path())
            self._handle_shutdown()

    # --- [REFACTOR] Slot Terpusat untuk Worker ---
    def _set_ui_for_conversion(self, is_running):
        self.start_btn.setEnabled(not is_running)
        self.stop_btn.setEnabled(is_running)
        self.category_list.setEnabled(not is_running)
        self.progress_bar.setValue(0)

    def _run_worker(self):
        self.thread = QThread()
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.conversion_finished.connect(self.on_conversion_finished)
        self.worker.conversion_error.connect(self.on_conversion_error)
        self.thread.start()
    
    def update_progress(self, value, text):
        self.progress_bar.setValue(value)
        # Untuk batch, tambahkan prefix info file
        if hasattr(self, 'current_mode_key') and self.batch_files[self.current_mode_key]:
             file_list = self.batch_files[self.current_mode_key]
             current_file_info = f"File {self.current_batch_index + 1}/{len(file_list)}"
             self.status_label.setText(f"({current_file_info}) {text}")
        else:
            self.status_label.setText(text)

    def on_conversion_finished(self, msg):
        self.progress_bar.setValue(100)
        is_batch_active = hasattr(self, 'current_mode_key') and self.batch_files[self.current_mode_key]
        if is_batch_active:
            self.current_batch_index += 1
            if self.thread: self.thread.quit(); self.thread.wait()
            self._start_next_batch_conversion()
        else:
            self.status_label.setText(msg)
            self._set_ui_for_conversion(False)
            if self.thread: self.thread.quit(); self.thread.wait()
            QMessageBox.information(self, self.lang["done"], msg)
            self._open_output_folder(self.get_active_output_path())
            self._handle_shutdown()

    def on_conversion_error(self, msg):
        self.status_label.setText(f"{self.lang['error_title']}: {msg}")
        self._set_ui_for_conversion(False)
        if self.thread: self.thread.quit(); self.thread.wait()
        QMessageBox.critical(self, self.lang["error_title"], msg)

    # --- Fungsi Bantuan Lainnya ---
    def get_active_output_path(self):
        page_index = self.stacked_widget.currentIndex()
        if page_index == 0: return self.vid_output_path_edit.text()
        if page_index == 1: return self.audio_output_path_edit.text()
        if page_index == 2: return self.va_output_path_edit.text()
        if page_index == 3: return self.img_output_path_edit.text()
        if page_index == 4: return self.output_path_edit.text()
        return ""

    def _handle_shutdown(self):
        if self.shutdown_checkbox.isChecked():
            reply = QMessageBox.question(self, self.lang["shutdown_confirm_title"], self.lang["shutdown_confirm_msg"],
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                if os.name == 'nt': os.system("shutdown /s /t 30")
                else: os.system("shutdown -h +1")

    def _open_output_folder(self, path):
        if not os.path.isdir(path): return
        try:
            if sys.platform == "win32": os.startfile(path)
            elif sys.platform == "darwin": subprocess.Popen(["open", path])
            else: subprocess.Popen(["xdg-open", path])
        except Exception as e: print(f"Gagal membuka folder output: {e}")

    def closeEvent(self, event):
        self.stop_conversion()
        event.accept()

# --- Entry Point Aplikasi ---
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RedesignedConverterApp()
    window.show()
    sys.exit(app.exec())