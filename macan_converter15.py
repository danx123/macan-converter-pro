import sys
import os
# [MODIFIKASI] Mengganti fitz dengan pypdfium2
import pypdfium2 as pdfium
import re
import subprocess
from PIL import Image
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QLineEdit, QFileDialog, QProgressBar,
    QComboBox, QMessageBox, QFrame, QTabWidget, QCheckBox, QGridLayout,
    QTextEdit
)
from PySide6.QtCore import QThread, QObject, Signal, Qt
from PySide6.QtGui import QIcon

# [MODIFIKASI] Mengimpor library baru yang dibutuhkan dan memperbarui pesan error
try:
    import docx
    from docx2pdf import convert
    import openpyxl
except ImportError:
    # Menampilkan pesan jika library penting tidak terinstal
    # Ini akan mencegah aplikasi crash saat startup jika dependensi kurang.
    print("PERINGATAN: Beberapa fungsionalitas mungkin tidak berfungsi. Pastikan Anda menginstal library yang dibutuhkan dengan 'pip install python-docx docx2pdf openpyxl pypdfium2'")


# --- [MODIFIKASI] Kamus untuk Teks Multi-Bahasa Diperbarui ---
LANGUAGES = {
    "id": {
        # Judul & Tab
        "window_title": "Macan Converter Pro",
        "main_title": "Macan Converter Pro",
        "tab_file": "File Converter",
        "tab_image": "Image Converter",
        "tab_video": "Video Converter",
        "tab_audio": "Audio Converter",
        "tab_video_audio": "Video-Audio Converter", # Baru
        "tab_about": "Tentang",

        # UI Umum
        "input_label": "1. Pilih File Input:",
        "output_label": "2. Pilih Folder Output:",
        "browse_btn": "Browse...",
        "start_conversion_btn": "Mulai Konversi",
        "ready_status": "Siap untuk mengonversi.",
        "batch_mode_checkbox": "Batch Mode (Konversi Banyak File Sekaligus)",

        # File Converter
        "input_placeholder_file": "Pilih file yang akan dikonversi...",
        "output_placeholder_folder": "Pilih folder untuk menyimpan hasil...",
        "conversion_mode_label": "3. Pilih Mode Konversi:",
        "mode_label": "Mode:",
        "output_format_label": "Format Output:",
        "ico_res_label": "Pilih Resolusi ICO:",
        # Mode baru ditambahkan
        "modes": ["PDF ke Gambar", "PNG ke ICO", "PDF ke TXT", "PDF ke DOCX", "PDF ke XLSX", "DOCX ke PDF", "Gambar ke PDF"],
        "start_file_conv_btn": "Mulai Konversi File",
        
        # Image Converter
        "input_img_label": "1. Pilih File Gambar Input:",
        "output_img_label": "2. Pilih Folder Output:",
        "input_placeholder_single_img": "Pilih satu file gambar...",
        "input_placeholder_multi_img": "Pilih satu atau beberapa file gambar...",
        "output_placeholder_img": "Pilih folder tujuan...",
        "img_options_label": "3. Atur Opsi Konversi Gambar:",
        "resolution_label": "Resolusi:",
        "quality_label": "Kualitas:",
        "img_formats": ["JPG", "PNG", "WEBP", "BMP", "GIF"],
        "img_qualities": ["Maksimum (100)", "Sangat Baik (95)", "Baik (85)", "Sedang (75)", "Rendah (50)"],
        "start_img_conv_btn": "Mulai Konversi Gambar",
        "ready_status_img": "Siap untuk mengonversi gambar.",

        # Video Converter
        "input_vid_label": "1. Pilih File Video Input:",
        "output_vid_label": "2. Pilih Folder Output:",
        "input_placeholder_single_vid": "Pilih satu file video...",
        "input_placeholder_multi_vid": "Pilih beberapa file video untuk konversi batch...",
        "output_placeholder_vid": "Pilih folder tujuan...",
        "vid_options_label": "3. Atur Opsi Konversi Video:",
        "vid_formats": ["mp4", "mkv", "avi", "mov", "webm", "gif"],
        "vid_qualities": ["Tinggi", "Sedang", "Rendah"],
        "start_vid_conv_btn": "Mulai Konversi Video",
        "ready_status_vid": "Siap untuk mengonversi video.",
        "shutdown_checkbox": "Matikan PC setelah selesai", # Baru
        "shutdown_confirm_title": "Konfirmasi Matikan PC", # Baru
        "shutdown_confirm_msg": "Konversi selesai. Apakah Anda yakin ingin mematikan PC sekarang?", # Baru

        # Audio Converter
        "input_audio_label": "1. Pilih File Audio Input:",
        "output_audio_label": "2. Pilih Folder Output:",
        "input_placeholder_single_audio": "Pilih satu file audio...",
        "input_placeholder_multi_audio": "Pilih beberapa file audio untuk konversi batch...",
        "output_placeholder_audio": "Pilih folder tujuan...",
        "audio_options_label": "3. Atur Opsi Konversi Audio:",
        "bitrate_label": "Bitrate:",
        "audio_formats": ["mp3", "wav", "aac", "flac", "ogg", "wma", "m4a"],
        "start_audio_conv_btn": "Mulai Konversi Audio",
        "ready_status_audio": "Siap untuk mengonversi audio.",

        # Video-Audio Converter (Baru)
        "input_va_label": "1. Pilih File Video Input:",
        "output_va_label": "2. Pilih Folder Output Audio:",
        "input_placeholder_single_va": "Pilih satu file video...",
        "input_placeholder_multi_va": "Pilih beberapa video untuk konversi batch...",
        "output_placeholder_va": "Pilih folder tujuan audio...",
        "va_options_label": "3. Atur Opsi Konversi Audio:",
        "start_va_conv_btn": "Mulai Ekstraksi Audio",
        "ready_status_va": "Siap untuk mengekstrak audio dari video.",
        
        # Pesan & Notifikasi
        "invalid_input_title": "Input Tidak Valid",
        "invalid_input_file_msg": "Silakan pilih file input yang valid.",
        "invalid_output_folder_msg": "Silakan pilih folder output yang valid.",
        "no_ico_size_msg": "Pilih setidaknya satu ukuran resolusi untuk file ICO.",
        "conversion_logic_error": "Logic konversi untuk mode '{mode}' tidak ditemukan.",
        "docx_conv_error_title": "Error Konversi DOCX",
        "docx_conv_error_msg": "Konversi DOCX ke PDF gagal. Pastikan Microsoft Word (Windows) atau LibreOffice (Linux/Mac) terinstal.",
        "ffmpeg_not_found_title": "FFmpeg Tidak Ditemukan",
        "ffmpeg_not_found_msg": "ffmpeg tidak ditemukan di direktori aplikasi atau di sistem PATH. Silakan letakkan file 'ffmpeg.exe' (Windows) atau 'ffmpeg' (Mac/Linux) di sebelah aplikasi ini.",
        "batch_no_files_msg": "Silakan pilih setidaknya satu file untuk mode batch.",
        "conversion_success_msg": "Sukses! Konversi telah selesai.",
        "batch_complete_title": "Batch Selesai",
        "batch_complete_msg": "Batch selesai! Semua {count} file telah berhasil dikonversi.",
        "error_title": "Error",
        "error_during_conversion": "Terjadi kesalahan: {error}",
        "error_in_batch": "Error pada file {index}: {error}\n\nProses batch dihentikan.",
        
        # Status Konversi
        "converting_page": "Mengonversi halaman {current} dari {total}...",
        "pdf_conversion_success": "Sukses! {total} halaman PDF telah dikonversi.",
        "opening_png": "Membuka file PNG...",
        "saving_ico": "Menyimpan ke format ICO dengan ukuran yang dipilih...",
        "done": "Selesai!",
        "png_to_ico_success": "Sukses! File PNG telah dikonversi ke ICO.",
        "opening_image": "Membuka {filename}...",
        "processing_image": "Memproses gambar...",
        "saving_image": "Gambar berhasil disimpan.",
        "image_conversion_success": "Sukses! Gambar telah dikonversi ke {format}.",
        "video_conversion_success": "Sukses! Video telah dikonversi ke {format}.",
        "getting_video_info": "Mendapatkan info video...",
        "converting_progress": "Mengonversi... {progress}%",
        "getting_audio_info": "Mendapatkan info audio...",
        "audio_conversion_success": "Sukses! Audio telah dikonversi ke {format}.",
        "preparing_conversion": "Mempersiapkan konversi untuk {filename}...",
        "converting_batch_file": "Mengonversi file {current} dari {total}: {filename}",
        
        # Tab About
        "about_content": """<b>Macan Converter Pro</b>
Versi 2.7.2

Dikembangkan oleh: Macan Angkasa
© 2025 Danx Exodus

Macan Converter Pro adalah aplikasi multi-konversi modern
untuk file, gambar, audio, dan video.
Mendukung batch mode, berbagai format populer, serta antarmuka
bersih dan mudah digunakan."""
    },
    "en": {
        # Titles & Tabs
        "window_title": "Macan Converter Pro",
        "main_title": "Macan Converter Pro",
        "tab_file": "File Converter",
        "tab_image": "Image Converter",
        "tab_video": "Video Converter",
        "tab_audio": "Audio Converter",
        "tab_video_audio": "Video-Audio Converter", # New
        "tab_about": "About",

        # Common UI
        "input_label": "1. Select Input File(s):",
        "output_label": "2. Select Output Folder:",
        "browse_btn": "Browse...",
        "start_conversion_btn": "Start Conversion",
        "ready_status": "Ready to convert.",
        "batch_mode_checkbox": "Batch Mode (Convert Multiple Files at Once)",

        # File Converter
        "input_placeholder_file": "Select a file to convert...",
        "output_placeholder_folder": "Select a folder to save the result...",
        "conversion_mode_label": "3. Select Conversion Mode:",
        "mode_label": "Mode:",
        "output_format_label": "Output Format:",
        "ico_res_label": "Select ICO Resolutions:",
        # New modes added
        "modes": ["PDF to Image", "PNG to ICO", "PDF to TXT", "PDF to DOCX", "PDF to XLSX", "DOCX to PDF", "Image to PDF"],
        "start_file_conv_btn": "Start File Conversion",
        
        # Image Converter
        "input_img_label": "1. Select Input Image File(s):",
        "output_img_label": "2. Select Output Folder:",
        "input_placeholder_single_img": "Select a single image file...",
        "input_placeholder_multi_img": "Select one or more image files...",
        "output_placeholder_img": "Select a destination folder...",
        "img_options_label": "3. Set Image Conversion Options:",
        "resolution_label": "Resolution:",
        "quality_label": "Quality:",
        "img_formats": ["JPEG", "PNG", "WEBP", "BMP", "GIF"],
        "img_qualities": ["Maximum (100)", "Very Good (95)", "Good (85)", "Medium (75)", "Low (50)"],
        "start_img_conv_btn": "Start Image Conversion",
        "ready_status_img": "Ready to convert images.",

        # Video Converter
        "input_vid_label": "1. Select Input Video File(s):",
        "output_vid_label": "2. Select Output Folder:",
        "input_placeholder_single_vid": "Select a single video file...",
        "input_placeholder_multi_vid": "Select multiple video files for batch conversion...",
        "output_placeholder_vid": "Select a destination folder...",
        "vid_options_label": "3. Set Video Conversion Options:",
        "vid_formats": ["mp4", "mkv", "avi", "mov", "webm", "gif"],
        "vid_qualities": ["High", "Medium", "Low"],
        "start_vid_conv_btn": "Start Video Conversion",
        "ready_status_vid": "Ready to convert videos.",
        "shutdown_checkbox": "Shutdown PC after completion", # New
        "shutdown_confirm_title": "Confirm PC Shutdown", # New
        "shutdown_confirm_msg": "Conversion is complete. Are you sure you want to shut down the PC now?", # New

        # Audio Converter
        "input_audio_label": "1. Select Input Audio File(s):",
        "output_audio_label": "2. Select Output Folder:",
        "input_placeholder_single_audio": "Select a single audio file...",
        "input_placeholder_multi_audio": "Select multiple audio files for batch conversion...",
        "output_placeholder_audio": "Select a destination folder...",
        "audio_options_label": "3. Set Audio Conversion Options:",
        "bitrate_label": "Bitrate:",
        "audio_formats": ["mp3", "wav", "aac", "flac", "ogg", "wma", "m4a"],
        "start_audio_conv_btn": "Start Audio Conversion",
        "ready_status_audio": "Ready to convert audio.",
        
        # Video-Audio Converter (New)
        "input_va_label": "1. Select Input Video File(s):",
        "output_va_label": "2. Select Audio Output Folder:",
        "input_placeholder_single_va": "Select a single video file...",
        "input_placeholder_multi_va": "Select multiple videos for batch conversion...",
        "output_placeholder_va": "Select audio destination folder...",
        "va_options_label": "3. Set Audio Conversion Options:",
        "start_va_conv_btn": "Start Audio Extraction",
        "ready_status_va": "Ready to extract audio from video.",
        
        # Messages & Notifications
        "invalid_input_title": "Invalid Input",
        "invalid_input_file_msg": "Please select a valid input file.",
        "invalid_output_folder_msg": "Please select a valid output folder.",
        "no_ico_size_msg": "Please select at least one resolution size for the ICO file.",
        "conversion_logic_error": "Conversion logic for mode '{mode}' not found.",
        "docx_conv_error_title": "DOCX Conversion Error",
        "docx_conv_error_msg": "DOCX to PDF conversion failed. Make sure Microsoft Word (Windows) or LibreOffice (Linux/Mac) is installed.",
        "ffmpeg_not_found_title": "FFmpeg Not Found",
        "ffmpeg_not_found_msg": "ffmpeg was not found in the application directory or in the system PATH. Please place 'ffmpeg.exe' (Windows) or 'ffmpeg' (Mac/Linux) next to this application.",
        "batch_no_files_msg": "Please select at least one file for batch mode.",
        "conversion_success_msg": "Success! The conversion is complete.",
        "batch_complete_title": "Batch Complete",
        "batch_complete_msg": "Batch finished! All {count} files have been successfully converted.",
        "error_title": "Error",
        "error_during_conversion": "An error occurred: {error}",
        "error_in_batch": "Error on file {index}: {error}\n\nBatch process stopped.",

        # Conversion Status
        "converting_page": "Converting page {current} of {total}...",
        "pdf_conversion_success": "Success! {total} PDF pages have been converted.",
        "opening_png": "Opening PNG file...",
        "saving_ico": "Saving to ICO format with selected sizes...",
        "done": "Done!",
        "png_to_ico_success": "Success! The PNG file has been converted to ICO.",
        "opening_image": "Opening {filename}...",
        "processing_image": "Processing image...",
        "saving_image": "Image saved successfully.",
        "image_conversion_success": "Success! The image has been converted to {format}.",
        "video_conversion_success": "Success! The video has been converted to {format}.",
        "getting_video_info": "Getting video info...",
        "converting_progress": "Converting... {progress}%",
        "getting_audio_info": "Getting audio info...",
        "audio_conversion_success": "Success! The audio has been converted to {format}.",
        "preparing_conversion": "Preparing conversion for {filename}...",
        "converting_batch_file": "Converting file {current} of {total}: {filename}",
        
        # About Tab
        "about_content": """<b>Macan Converter Pro</b>
Version 2.7.0

Developed by: Macan Angkasa
© 2025 Danx Exodus

Macan Converter Pro is a modern multi-conversion application
for files, images, audio, and video.
It supports batch mode, various popular formats, and features a
clean and easy-to-use interface."""
    }
}


# --- [MODIFIKASI] Kelas Worker untuk Proses Konversi File diperbarui ---
class FileConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    # `input_path` sekarang bisa berupa string atau list
    def __init__(self, input_path, output_path, mode, out_format, ico_sizes=None, lang_dict=None):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.mode = mode
        self.out_format = out_format
        self.ico_sizes = ico_sizes if ico_sizes else []
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def run(self):
        try:
            mode_map = {
                "PDF ke Gambar": self._convert_pdf_to_image, "PDF to Image": self._convert_pdf_to_image,
                "PNG ke ICO": self._convert_png_to_ico, "PNG to ICO": self._convert_png_to_ico,
                "PDF ke TXT": self._convert_pdf_to_txt, "PDF to TXT": self._convert_pdf_to_txt,
                "PDF ke DOCX": self._convert_pdf_to_docx, "PDF to DOCX": self._convert_pdf_to_docx,
                "PDF ke XLSX": self._convert_pdf_to_xlsx, "PDF to XLSX": self._convert_pdf_to_xlsx,
                "DOCX ke PDF": self._convert_docx_to_pdf, "DOCX to PDF": self._convert_docx_to_pdf,
                "Gambar ke PDF": self._convert_image_to_pdf, "Image to PDF": self._convert_image_to_pdf,
            }
            conversion_func = mode_map.get(self.mode)
            if conversion_func:
                conversion_func()
            else:
                self.conversion_error.emit(self.lang["conversion_logic_error"].format(mode=self.mode))
        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False

    # [MODIFIKASI] Diubah untuk menggunakan PyPDFium2
    def _convert_pdf_to_image(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]

        for i in range(total_pages):
            if not self.is_running: break
            page = pdf[i]
            # scale=200/72 karena DPI default adalah 72. Ini setara dengan dpi=200 di PyMuPDF
            image = page.render(scale=200/72).to_pil()
            output_filename = os.path.join(self.output_path, f"{base_name}_page_{i+1}.{self.out_format.lower()}")
            image.save(output_filename)
            progress = int(((i + 1) / total_pages) * 100)
            status_text = self.lang["converting_page"].format(current=i+1, total=total_pages)
            self.progress_updated.emit(progress, status_text)
            
        if self.is_running:
            self.conversion_finished.emit(self.lang["pdf_conversion_success"].format(total=total_pages))

    def _convert_png_to_ico(self):
        self.progress_updated.emit(20, self.lang["opening_png"])
        if not self.ico_sizes:
            self.conversion_error.emit(self.lang["no_ico_size_msg"])
            return
            
        img = Image.open(self.input_path)
        self.progress_updated.emit(60, self.lang["saving_ico"])
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.ico")
        img.save(output_filename, format='ICO', sizes=self.ico_sizes)
        self.progress_updated.emit(100, self.lang["done"])
        self.conversion_finished.emit(self.lang["png_to_ico_success"])
    
    # [MODIFIKASI] Diubah untuk menggunakan PyPDFium2
    def _convert_pdf_to_txt(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.txt")
        
        full_text = ""
        for i in range(total_pages):
            if not self.is_running: break
            page = pdf[i]
            text_page = page.get_textpage()
            full_text += text_page.get_text_range() + "\n\n"
            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))
        
        with open(output_filename, "w", encoding="utf-8") as f:
            f.write(full_text)
            
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])
    
    # [MODIFIKASI] Diubah untuk menggunakan PyPDFium2
    def _convert_pdf_to_docx(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.docx")
        
        word_doc = docx.Document()
        for i in range(total_pages):
            if not self.is_running: break
            page = pdf[i]
            text_page = page.get_textpage()
            word_doc.add_paragraph(text_page.get_text_range())
            if i < total_pages - 1:
                word_doc.add_page_break()
            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))

        word_doc.save(output_filename)
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

    # [MODIFIKASI] Diubah untuk menggunakan PyPDFium2
    def _convert_pdf_to_xlsx(self):
        pdf = pdfium.PdfDocument(self.input_path)
        total_pages = len(pdf)
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.xlsx")
        
        workbook = openpyxl.Workbook()
        if "Sheet" in workbook.sheetnames:
            workbook.remove(workbook["Sheet"])

        for i in range(total_pages):
            if not self.is_running: break
            sheet = workbook.create_sheet(title=f"Page_{i+1}")
            
            page = pdf[i]
            text_page = page.get_textpage()
            page_text = text_page.get_text_range()
            
            # PyPDFium2 tidak memiliki deteksi tabel bawaan seperti PyMuPDF.
            # Jadi, kita akan mengekstrak teks dan membaginya per baris sebagai fallback.
            lines = page_text.splitlines()
            for r, line in enumerate(lines, 1):
                sheet.cell(row=r, column=1, value=line)

            progress = int(((i + 1) / total_pages) * 100)
            self.progress_updated.emit(progress, self.lang["converting_page"].format(current=i+1, total=total_pages))
        
        workbook.save(output_filename)
        if self.is_running:
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

    def _convert_docx_to_pdf(self):
        self.progress_updated.emit(10, "Mempersiapkan konversi DOCX...")
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.pdf")
        
        try:
            self.progress_updated.emit(50, "Mengonversi file...")
            convert(self.input_path, output_filename)
            self.progress_updated.emit(100, self.lang["done"])
            self.conversion_finished.emit(self.lang["conversion_success_msg"])
        except Exception as e:
            # Gagal karena Word/LibreOffice tidak terinstal
            self.conversion_error.emit(f"{self.lang['docx_conv_error_msg']} Detail: {e}")

    # [MODIFIKASI] Diubah untuk menggunakan Pillow (tanpa fitz atau pypdfium2)
    def _convert_image_to_pdf(self):
        total_files = len(self.input_path)
        if total_files == 0:
            self.conversion_error.emit(self.lang["batch_no_files_msg"])
            return

        base_name = os.path.splitext(os.path.basename(self.input_path[0]))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}_converted.pdf")
        
        images_pil = []
        for i, img_path in enumerate(self.input_path):
            if not self.is_running: break
            try:
                img = Image.open(img_path).convert("RGB")
                images_pil.append(img)
            except Exception as e:
                self.conversion_error.emit(f"Gagal membuka gambar {os.path.basename(img_path)}: {e}")
                return

            progress = int(((i + 1) / total_files) * 100)
            status_text = f"Memproses gambar {i+1} dari {total_files}..."
            self.progress_updated.emit(progress, status_text)
        
        if self.is_running and images_pil:
            first_image = images_pil[0]
            other_images = images_pil[1:]
            first_image.save(output_filename, "PDF" ,resolution=100.0, save_all=True, append_images=other_images)
            self.conversion_finished.emit(self.lang["conversion_success_msg"])

# --- Kelas Worker untuk Proses Konversi Gambar di Latar Belakang ---
class ImageConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, input_path, output_path, out_format, resolution, quality_str, lang_dict=None):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.out_format = out_format
        self.resolution = resolution
        self.quality_str = quality_str
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def stop(self):
        self.is_running = False
        
    def run(self):
        try:
            self.progress_updated.emit(0, self.lang["opening_image"].format(filename=os.path.basename(self.input_path)))
            if not self.is_running: return
            
            img = Image.open(self.input_path)

            if self.resolution != "Original Size":
                try:
                    res_parts = self.resolution.split(' ')[0].split('x')
                    new_size = (int(res_parts[0]), int(res_parts[1]))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                except (ValueError, IndexError):
                    self.conversion_error.emit(f"Invalid resolution format: {self.resolution}")
                    return

            self.progress_updated.emit(50, self.lang["processing_image"])
            if not self.is_running: return

            base_name = os.path.splitext(os.path.basename(self.input_path))[0]
            output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format.lower()}")
            
            save_options = {}
            if self.out_format.lower() in ['jpeg', 'jpg', 'webp']:
                quality_map = {
                    "Maksimum (100)": 100, "Sangat Baik (95)": 95, "Baik (85)": 85,
                    "Sedang (75)": 75, "Rendah (50)": 50,
                    "Maximum (100)": 100, "Very Good (95)": 95, "Good (85)": 85,
                    "Medium (75)": 75, "Low (50)": 50
                }
                save_options['quality'] = quality_map.get(self.quality_str, 85)
            
            if self.out_format.lower() in ['jpeg', 'jpg']:
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')

            img.save(output_filename, **save_options)
            self.progress_updated.emit(100, self.lang["saving_image"])
            self.conversion_finished.emit(self.lang["image_conversion_success"].format(format=self.out_format.upper()))

        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))


# --- Kelas Worker untuk Proses Konversi Video di Latar Belakang ---
class VideoConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, ffmpeg_path, input_path, output_path, out_format, resolution, quality, lang_dict=None):
        super().__init__()
        self.ffmpeg_path = ffmpeg_path
        self.input_path = input_path
        self.output_path = output_path
        self.out_format = out_format
        self.resolution = resolution
        self.quality = quality
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def _get_media_duration(self):
        command = [self.ffmpeg_path, '-i', self.input_path]
        try:
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            result = subprocess.run(command, capture_output=True, text=True, stderr=subprocess.STDOUT, startupinfo=startupinfo)
            output = result.stdout
            duration_search = re.search(r"Duration: (\d{2}):(\d{2}):(\d{2})\.(\d{2})", output)
            if duration_search:
                hours, minutes, seconds = map(int, duration_search.groups()[:3])
                return (hours * 3600) + (minutes * 60) + seconds
        except Exception:
            return None
        return None

    def run(self):
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format}")
        command = [self.ffmpeg_path, '-i', self.input_path]
        
        quality_map = {"Tinggi": "18", "Sedang": "23", "Rendah": "28", "High": "18", "Medium": "23", "Low": "28"}
        command.extend(['-crf', quality_map.get(self.quality, "23")])

        resolution_map = {"360p": "360", "480p": "480", "720p": "720", "1080p": "1080", "2K": "1440", "4K": "2160"}
        if self.resolution in resolution_map:
            command.extend(['-vf', f'scale=-2:{resolution_map[self.resolution]}'])

        command.extend(['-y', output_filename])
        
        try:
            self.progress_updated.emit(0, self.lang["getting_video_info"])
            total_duration = self._get_media_duration()
            
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, universal_newlines=True, encoding='utf-8', startupinfo=startupinfo)

            for line in iter(process.stdout.readline, ""):
                if not self.is_running:
                    process.terminate()
                    break
                if total_duration:
                    time_search = re.search(r"time=(\d{2}):(\d{2}):(\d{2})\.(\d{2})", line)
                    if time_search:
                        hours, minutes, seconds = map(int, time_search.groups()[:3])
                        current_time = (hours * 3600) + (minutes * 60) + seconds
                        progress = int((current_time / total_duration) * 100)
                        self.progress_updated.emit(progress, self.lang["converting_progress"].format(progress=progress))

            process.wait()

            if self.is_running and process.returncode == 0:
                self.conversion_finished.emit(self.lang["video_conversion_success"].format(format=self.out_format.upper()))
            elif self.is_running:
                self.conversion_error.emit(f"Failed to convert {os.path.basename(self.input_path)}. Exit code: {process.returncode}")

        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False

# --- Kelas Worker untuk Proses Konversi Audio di Latar Belakang ---
class AudioConversionWorker(QObject):
    progress_updated = Signal(int, str)
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    def __init__(self, ffmpeg_path, input_path, output_path, out_format, bitrate, lang_dict=None):
        super().__init__()
        self.ffmpeg_path = ffmpeg_path
        self.input_path = input_path
        self.output_path = output_path
        self.out_format = out_format
        self.bitrate = bitrate
        self.is_running = True
        self.lang = lang_dict if lang_dict else LANGUAGES["id"]

    def _get_media_duration(self):
        command = [self.ffmpeg_path, '-i', self.input_path]
        try:
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            result = subprocess.run(command, capture_output=True, text=True, stderr=subprocess.STDOUT, startupinfo=startupinfo)
            output = result.stdout
            duration_search = re.search(r"Duration: (\d{2}):(\d{2}):(\d{2})\.(\d{2})", output)
            if duration_search:
                hours, minutes, seconds = map(int, duration_search.groups()[:3])
                return (hours * 3600) + (minutes * 60) + seconds
        except Exception:
            return None
        return None

    def run(self):
        base_name = os.path.splitext(os.path.basename(self.input_path))[0]
        output_filename = os.path.join(self.output_path, f"{base_name}.{self.out_format}")
        command = [self.ffmpeg_path, '-i', self.input_path, '-vn'] # -vn adalah kunci untuk menghapus video
        
        command.extend(['-b:a', self.bitrate.split(' ')[0] + 'k'])
        command.extend(['-y', output_filename])
        
        try:
            self.progress_updated.emit(0, self.lang["getting_audio_info"])
            total_duration = self._get_media_duration()
            
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, universal_newlines=True, encoding='utf-8', startupinfo=startupinfo)

            for line in iter(process.stdout.readline, ""):
                if not self.is_running:
                    process.terminate()
                    break
                if total_duration:
                    time_search = re.search(r"time=(\d{2}):(\d{2}):(\d{2})\.(\d{2})", line)
                    if time_search:
                        hours, minutes, seconds = map(int, time_search.groups()[:3])
                        current_time = (hours * 3600) + (minutes * 60) + seconds
                        progress = int((current_time / total_duration) * 100)
                        self.progress_updated.emit(progress, self.lang["converting_progress"].format(progress=progress))

            process.wait()

            if self.is_running and process.returncode == 0:
                self.conversion_finished.emit(self.lang["audio_conversion_success"].format(format=self.out_format.upper()))
            elif self.is_running:
                self.conversion_error.emit(f"Failed to convert {os.path.basename(self.input_path)}. Exit code: {process.returncode}")

        except Exception as e:
            self.conversion_error.emit(self.lang["error_during_conversion"].format(error=str(e)))

    def stop(self):
        self.is_running = False

# --- Kelas Utama Aplikasi (UI) ---
class ConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_lang = "id"  # Bahasa default
        self.lang = LANGUAGES[self.current_lang]
        
        self.setWindowTitle(self.lang["window_title"])
        self.setGeometry(100, 100, 650, 600)
        icon_path = "icon.ico"
        if hasattr(sys, "_MEIPASS"):
            icon_path = os.path.join(sys._MEIPASS, icon_path)
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        self.thread = None
        self.worker = None
        
        # [MODIFIKASI] Menambahkan list batch untuk tab baru
        self.file_batch_files = []
        self.image_batch_files = []
        self.current_image_batch_index = 0
        self.video_batch_files = []
        self.current_batch_index = 0
        self.audio_batch_files = []
        self.current_audio_batch_index = 0
        self.va_batch_files = [] # Batch untuk Video-Audio
        self.current_va_batch_index = 0
        
        self._setup_ui()
        self._apply_stylesheet()
        self._update_ui_for_mode(self.mode_combo.currentText())

    def _setup_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)

        header_layout = QHBoxLayout()
        self.title_label = QLabel(self.lang["main_title"])
        self.title_label.setObjectName("titleLabel")
        
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["Indonesia", "English"])
        self.lang_combo.currentIndexChanged.connect(self.change_language)
        
        header_layout.addWidget(self.title_label, 1)
        header_layout.addStretch()
        header_layout.addWidget(QLabel("Bahasa/Language:"))
        header_layout.addWidget(self.lang_combo)
        main_layout.addLayout(header_layout)

        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)

        self.file_converter_tab = self._create_file_converter_tab()
        self.image_converter_tab = self._create_image_converter_tab()
        self.video_converter_tab = self._create_video_converter_tab()
        self.audio_converter_tab = self._create_audio_converter_tab()
        self.video_audio_tab = self._create_video_audio_converter_tab() # [MODIFIKASI]
        self.about_tab = self._create_about_tab()
        
        self.tab_widget.addTab(self.file_converter_tab, self.lang["tab_file"])
        self.tab_widget.addTab(self.image_converter_tab, self.lang["tab_image"])
        self.tab_widget.addTab(self.video_converter_tab, self.lang["tab_video"])
        self.tab_widget.addTab(self.audio_converter_tab, self.lang["tab_audio"])
        self.tab_widget.addTab(self.video_audio_tab, self.lang["tab_video_audio"]) # [MODIFIKASI]
        self.tab_widget.addTab(self.about_tab, self.lang["tab_about"])

    def change_language(self):
        selected_lang = self.lang_combo.currentText()
        self.current_lang = "en" if selected_lang == "English" else "id"
        self.lang = LANGUAGES[self.current_lang]
        self._retranslate_ui()

    def _retranslate_ui(self):
        """Menerjemahkan semua teks UI ke bahasa yang dipilih."""
        self.setWindowTitle(self.lang["window_title"])
        self.title_label.setText(self.lang["main_title"])

        # Nama Tab
        self.tab_widget.setTabText(0, self.lang["tab_file"])
        self.tab_widget.setTabText(1, self.lang["tab_image"])
        self.tab_widget.setTabText(2, self.lang["tab_video"])
        self.tab_widget.setTabText(3, self.lang["tab_audio"])
        self.tab_widget.setTabText(4, self.lang["tab_video_audio"]) # [MODIFIKASI]
        self.tab_widget.setTabText(5, self.lang["tab_about"])     # [MODIFIKASI]

        # Tab: File Converter
        self.file_input_label.setText(self.lang["input_label"])
        self.input_path_edit.setPlaceholderText(self.lang["input_placeholder_file"])
        self.browse_input_btn.setText(self.lang["browse_btn"])
        self.file_output_label.setText(self.lang["output_label"])
        self.output_path_edit.setPlaceholderText(self.lang["output_placeholder_folder"])
        self.browse_output_btn.setText(self.lang["browse_btn"])
        self.conv_mode_label.setText(self.lang["conversion_mode_label"])
        self.mode_label.setText(self.lang["mode_label"])
        self.mode_combo.clear()
        self.mode_combo.addItems(self.lang["modes"])
        self.format_label.setText(self.lang["output_format_label"])
        self.ico_res_label.setText(self.lang["ico_res_label"])
        self.convert_btn.setText(self.lang["start_file_conv_btn"])
        self.status_label.setText(self.lang["ready_status"])

        # Tab: Image Converter
        self.img_input_label.setText(self.lang["input_img_label"])
        self._update_image_input_ui()
        self.img_browse_btn.setText(self.lang["browse_btn"])
        self.img_output_label.setText(self.lang["output_img_label"])
        self.img_output_path_edit.setPlaceholderText(self.lang["output_placeholder_img"])
        self.img_browse_output_btn.setText(self.lang["browse_btn"])
        self.img_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.img_options_label.setText(self.lang["img_options_label"])
        self.img_format_label.setText(self.lang["output_format_label"])
        self.img_format_combo.clear()
        self.img_format_combo.addItems(self.lang["img_formats"])
        self.img_res_label.setText(self.lang["resolution_label"])
        self.img_quality_label.setText(self.lang["quality_label"])
        self.img_quality_combo.clear()
        self.img_quality_combo.addItems(self.lang["img_qualities"])
        self.img_quality_combo.setCurrentText(self.lang["img_qualities"][2])
        self.img_convert_btn.setText(self.lang["start_img_conv_btn"])
        self.img_status_label.setText(self.lang["ready_status_img"])

        # Tab: Video Converter
        self.vid_input_label.setText(self.lang["input_vid_label"])
        self._update_video_input_ui()
        self.vid_browse_btn.setText(self.lang["browse_btn"])
        self.vid_output_label.setText(self.lang["output_vid_label"])
        self.vid_output_path_edit.setPlaceholderText(self.lang["output_placeholder_vid"])
        self.vid_browse_output_btn.setText(self.lang["browse_btn"])
        self.batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.vid_options_label.setText(self.lang["vid_options_label"])
        self.vid_format_label.setText(self.lang["output_format_label"])
        self.vid_format_combo.clear()
        self.vid_format_combo.addItems(self.lang["vid_formats"])
        self.vid_res_label.setText(self.lang["resolution_label"])
        self.vid_quality_label.setText(self.lang["quality_label"])
        self.vid_quality_combo.clear()
        self.vid_quality_combo.addItems(self.lang["vid_qualities"])
        self.vid_quality_combo.setCurrentText(self.lang["vid_qualities"][1])
        self.vid_convert_btn.setText(self.lang["start_vid_conv_btn"])
        self.vid_status_label.setText(self.lang["ready_status_vid"])
        self.shutdown_checkbox.setText(self.lang["shutdown_checkbox"]) # [MODIFIKASI]

        # Tab: Audio Converter
        self.audio_input_label.setText(self.lang["input_audio_label"])
        self._update_audio_input_ui()
        self.audio_browse_btn.setText(self.lang["browse_btn"])
        self.audio_output_label.setText(self.lang["output_audio_label"])
        self.audio_output_path_edit.setPlaceholderText(self.lang["output_placeholder_audio"])
        self.audio_browse_output_btn.setText(self.lang["browse_btn"])
        self.audio_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.audio_options_label.setText(self.lang["audio_options_label"])
        self.audio_format_label.setText(self.lang["output_format_label"])
        self.audio_format_combo.clear()
        self.audio_format_combo.addItems(self.lang["audio_formats"])
        self.audio_bitrate_label.setText(self.lang["bitrate_label"])
        self.audio_convert_btn.setText(self.lang["start_audio_conv_btn"])
        self.audio_status_label.setText(self.lang["ready_status_audio"])
        
        # [MODIFIKASI] Tab: Video-Audio Converter
        self.va_input_label.setText(self.lang["input_va_label"])
        self._update_va_input_ui()
        self.va_browse_btn.setText(self.lang["browse_btn"])
        self.va_output_label.setText(self.lang["output_va_label"])
        self.va_output_path_edit.setPlaceholderText(self.lang["output_placeholder_va"])
        self.va_browse_output_btn.setText(self.lang["browse_btn"])
        self.va_batch_mode_checkbox.setText(self.lang["batch_mode_checkbox"])
        self.va_options_label.setText(self.lang["va_options_label"])
        self.va_format_label.setText(self.lang["output_format_label"])
        self.va_format_combo.clear()
        self.va_format_combo.addItems(self.lang["audio_formats"])
        self.va_bitrate_label.setText(self.lang["bitrate_label"])
        self.va_convert_btn.setText(self.lang["start_va_conv_btn"])
        self.va_status_label.setText(self.lang["ready_status_va"])

        # Tab: About
        self.about_text.setHtml(self.lang["about_content"])

    def _create_about_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(20, 20, 20, 20)
        self.about_text = QTextEdit()
        self.about_text.setReadOnly(True)
        self.about_text.setHtml(self.lang["about_content"])
        self.about_text.setObjectName("aboutText")
        layout.addWidget(self.about_text)
        return tab_widget

    def _create_file_converter_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(10, 20, 10, 10)
        layout.setSpacing(15)

        io_frame = QFrame()
        io_frame.setObjectName("ioFrame")
        io_layout = QVBoxLayout(io_frame)
        io_layout.setSpacing(10)
        
        self.file_input_label = QLabel(self.lang["input_label"])
        io_layout.addWidget(self.file_input_label)
        input_layout = QHBoxLayout()
        self.input_path_edit = QLineEdit()
        self.input_path_edit.setPlaceholderText(self.lang["input_placeholder_file"])
        self.input_path_edit.setReadOnly(True)
        self.browse_input_btn = QPushButton(self.lang["browse_btn"])
        self.browse_input_btn.clicked.connect(self.browse_input_file)
        input_layout.addWidget(self.input_path_edit)
        input_layout.addWidget(self.browse_input_btn)
        io_layout.addLayout(input_layout)

        self.file_output_label = QLabel(self.lang["output_label"])
        io_layout.addWidget(self.file_output_label)
        output_layout = QHBoxLayout()
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setPlaceholderText(self.lang["output_placeholder_folder"])
        self.output_path_edit.setReadOnly(True)
        self.browse_output_btn = QPushButton(self.lang["browse_btn"])
        self.browse_output_btn.clicked.connect(self.browse_output_folder)
        output_layout.addWidget(self.output_path_edit)
        output_layout.addWidget(self.browse_output_btn)
        io_layout.addLayout(output_layout)
        layout.addWidget(io_frame)

        settings_frame = QFrame()
        settings_frame.setObjectName("ioFrame")
        settings_layout = QVBoxLayout(settings_frame)
        self.conv_mode_label = QLabel(self.lang["conversion_mode_label"])
        settings_layout.addWidget(self.conv_mode_label)
        
        settings_grid_layout = QHBoxLayout()
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(self.lang["modes"])
        self.mode_combo.currentTextChanged.connect(self._update_ui_for_mode)
        
        self.format_combo = QComboBox()
        self.format_combo.addItems(["PNG", "JPG", "WEBP"])
        
        self.mode_label = QLabel(self.lang["mode_label"])
        settings_grid_layout.addWidget(self.mode_label)
        settings_grid_layout.addWidget(self.mode_combo, 1)
        settings_grid_layout.addSpacing(20)
        self.format_label = QLabel(self.lang["output_format_label"])
        settings_grid_layout.addWidget(self.format_label)
        settings_grid_layout.addWidget(self.format_combo, 1)
        settings_layout.addLayout(settings_grid_layout)
        
        self.ico_sizes_frame = QFrame()
        self.ico_sizes_frame.setObjectName("icoFrame")
        ico_layout = QVBoxLayout(self.ico_sizes_frame)
        self.ico_res_label = QLabel(self.lang["ico_res_label"])
        ico_layout.addWidget(self.ico_res_label)
        ico_grid = QGridLayout()
        self.ico_checkboxes = {}
        sizes = ["16x16", "32x32", "48x48", "64x64", "128x128", "256x256"]
        positions = [(i, j) for i in range(2) for j in range(3)]
        for position, size in zip(positions, sizes):
            checkbox = QCheckBox(size)
            checkbox.setChecked(True)
            self.ico_checkboxes[size] = checkbox
            ico_grid.addWidget(checkbox, *position)
        ico_layout.addLayout(ico_grid)
        settings_layout.addWidget(self.ico_sizes_frame)
        layout.addWidget(settings_frame)

        self.convert_btn = QPushButton(self.lang["start_file_conv_btn"])
        self.convert_btn.setObjectName("convertButton")
        self.convert_btn.clicked.connect(self.start_file_conversion)
        layout.addWidget(self.convert_btn)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)
        layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel(self.lang["ready_status"])
        self.status_label.setObjectName("statusLabel")
        layout.addWidget(self.status_label)

        layout.addStretch()
        return tab_widget

    def _create_image_converter_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(10, 20, 10, 10)
        layout.setSpacing(15)

        img_io_frame = QFrame()
        img_io_frame.setObjectName("ioFrame")
        img_io_layout = QVBoxLayout(img_io_frame)
        
        self.img_input_label = QLabel(self.lang["input_img_label"])
        img_io_layout.addWidget(self.img_input_label)
        img_input_layout = QHBoxLayout()
        self.img_input_path_edit = QLineEdit()
        self.img_input_path_edit.setReadOnly(True)
        self.img_browse_btn = QPushButton(self.lang["browse_btn"])
        self.img_browse_btn.clicked.connect(self.browse_image_input_file)
        img_input_layout.addWidget(self.img_input_path_edit)
        img_input_layout.addWidget(self.img_browse_btn)
        img_io_layout.addLayout(img_input_layout)

        self.img_output_label = QLabel(self.lang["output_img_label"])
        img_io_layout.addWidget(self.img_output_label)
        img_output_layout = QHBoxLayout()
        self.img_output_path_edit = QLineEdit()
        self.img_output_path_edit.setPlaceholderText(self.lang["output_placeholder_img"])
        self.img_output_path_edit.setReadOnly(True)
        self.img_browse_output_btn = QPushButton(self.lang["browse_btn"])
        self.img_browse_output_btn.clicked.connect(self.browse_image_output_folder)
        img_output_layout.addWidget(self.img_output_path_edit)
        img_output_layout.addWidget(self.img_browse_output_btn)
        img_io_layout.addLayout(img_output_layout)

        self.img_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.img_batch_mode_checkbox.stateChanged.connect(self._update_image_input_ui)
        img_io_layout.addWidget(self.img_batch_mode_checkbox)
        
        layout.addWidget(img_io_frame)

        img_settings_frame = QFrame()
        img_settings_frame.setObjectName("ioFrame")
        img_settings_layout = QVBoxLayout(img_settings_frame)
        self.img_options_label = QLabel(self.lang["img_options_label"])
        img_settings_layout.addWidget(self.img_options_label)

        format_layout = QHBoxLayout()
        self.img_format_label = QLabel(self.lang["output_format_label"])
        format_layout.addWidget(self.img_format_label)
        self.img_format_combo = QComboBox()
        self.img_format_combo.addItems(self.lang["img_formats"])
        self.img_format_combo.currentTextChanged.connect(self._update_image_options)
        format_layout.addWidget(self.img_format_combo, 1)
        img_settings_layout.addLayout(format_layout)

        res_quality_layout = QHBoxLayout()
        
        self.img_res_label = QLabel(self.lang["resolution_label"])
        res_quality_layout.addWidget(self.img_res_label)
        self.img_resolution_combo = QComboBox()
        self.img_resolution_combo.addItems([
            "Original Size", "320x240", "640x480", "800x600", "1280x720 (HD)", "1920x1080 (Full HD)", "2560x1440 (2K)", "3840x2160 (4K)"
        ])
        res_quality_layout.addWidget(self.img_resolution_combo, 1)
        
        res_quality_layout.addSpacing(20)

        self.img_quality_label = QLabel(self.lang["quality_label"])
        res_quality_layout.addWidget(self.img_quality_label)
        self.img_quality_combo = QComboBox()
        self.img_quality_combo.addItems(self.lang["img_qualities"])
        self.img_quality_combo.setCurrentText(self.lang["img_qualities"][2]) 
        res_quality_layout.addWidget(self.img_quality_combo, 1)

        img_settings_layout.addLayout(res_quality_layout)
        layout.addWidget(img_settings_frame)

        self.img_convert_btn = QPushButton(self.lang["start_img_conv_btn"])
        self.img_convert_btn.setObjectName("convertButton")
        self.img_convert_btn.clicked.connect(self.start_image_conversion)
        layout.addWidget(self.img_convert_btn)

        self.img_progress_bar = QProgressBar()
        self.img_progress_bar.setValue(0)
        self.img_progress_bar.setTextVisible(False)
        layout.addWidget(self.img_progress_bar)
        
        self.img_status_label = QLabel(self.lang["ready_status_img"])
        self.img_status_label.setObjectName("statusLabel")
        layout.addWidget(self.img_status_label)

        layout.addStretch()
        self._update_image_input_ui()
        self._update_image_options()
        return tab_widget

    def _create_video_converter_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(10, 20, 10, 10)
        layout.setSpacing(15)

        video_io_frame = QFrame()
        video_io_frame.setObjectName("ioFrame")
        video_io_layout = QVBoxLayout(video_io_frame)
        
        self.vid_input_label = QLabel(self.lang["input_vid_label"])
        video_io_layout.addWidget(self.vid_input_label)
        vid_input_layout = QHBoxLayout()
        self.vid_input_path_edit = QLineEdit()
        self.vid_input_path_edit.setReadOnly(True)
        self.vid_browse_btn = QPushButton(self.lang["browse_btn"])
        self.vid_browse_btn.clicked.connect(self.browse_video_input_file)
        vid_input_layout.addWidget(self.vid_input_path_edit)
        vid_input_layout.addWidget(self.vid_browse_btn)
        video_io_layout.addLayout(vid_input_layout)

        self.vid_output_label = QLabel(self.lang["output_vid_label"])
        video_io_layout.addWidget(self.vid_output_label)
        vid_output_layout = QHBoxLayout()
        self.vid_output_path_edit = QLineEdit()
        self.vid_output_path_edit.setPlaceholderText(self.lang["output_placeholder_vid"])
        self.vid_output_path_edit.setReadOnly(True)
        self.vid_browse_output_btn = QPushButton(self.lang["browse_btn"])
        self.vid_browse_output_btn.clicked.connect(self.browse_video_output_folder)
        vid_output_layout.addWidget(self.vid_output_path_edit)
        vid_output_layout.addWidget(self.vid_browse_output_btn)
        video_io_layout.addLayout(vid_output_layout)

        self.batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.batch_mode_checkbox.stateChanged.connect(self._update_video_input_ui)
        video_io_layout.addWidget(self.batch_mode_checkbox)
        
        layout.addWidget(video_io_frame)

        vid_settings_frame = QFrame()
        vid_settings_frame.setObjectName("ioFrame")
        vid_settings_layout = QVBoxLayout(vid_settings_frame)
        self.vid_options_label = QLabel(self.lang["vid_options_label"])
        vid_settings_layout.addWidget(self.vid_options_label)

        format_layout = QHBoxLayout()
        self.vid_format_label = QLabel(self.lang["output_format_label"])
        format_layout.addWidget(self.vid_format_label)
        self.vid_format_combo = QComboBox()
        self.vid_format_combo.addItems(self.lang["vid_formats"])
        format_layout.addWidget(self.vid_format_combo, 1)
        vid_settings_layout.addLayout(format_layout)

        res_quality_layout = QHBoxLayout()
        self.vid_res_label = QLabel(self.lang["resolution_label"])
        res_quality_layout.addWidget(self.vid_res_label)
        self.vid_resolution_combo = QComboBox()
        self.vid_resolution_combo.addItems(["Original Size", "360p", "480p", "720p", "1080p", "2K", "4K"])
        res_quality_layout.addWidget(self.vid_resolution_combo, 1)
        res_quality_layout.addSpacing(20)
        self.vid_quality_label = QLabel(self.lang["quality_label"])
        res_quality_layout.addWidget(self.vid_quality_label)
        self.vid_quality_combo = QComboBox()
        self.vid_quality_combo.addItems(self.lang["vid_qualities"])
        self.vid_quality_combo.setCurrentText(self.lang["vid_qualities"][1]) 
        res_quality_layout.addWidget(self.vid_quality_combo, 1)
        vid_settings_layout.addLayout(res_quality_layout)
        layout.addWidget(vid_settings_frame)
        
        # [MODIFIKASI] Checkbox untuk mematikan PC
        self.shutdown_checkbox = QCheckBox(self.lang["shutdown_checkbox"])
        layout.addWidget(self.shutdown_checkbox)

        self.vid_convert_btn = QPushButton(self.lang["start_vid_conv_btn"])
        self.vid_convert_btn.setObjectName("convertButton")
        self.vid_convert_btn.clicked.connect(self.start_video_conversion)
        layout.addWidget(self.vid_convert_btn)

        self.vid_progress_bar = QProgressBar()
        self.vid_progress_bar.setValue(0)
        self.vid_progress_bar.setTextVisible(False)
        layout.addWidget(self.vid_progress_bar)
        
        self.vid_status_label = QLabel(self.lang["ready_status_vid"])
        self.vid_status_label.setObjectName("statusLabel")
        layout.addWidget(self.vid_status_label)

        layout.addStretch()
        self._update_video_input_ui()
        return tab_widget

    def _create_audio_converter_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(10, 20, 10, 10)
        layout.setSpacing(15)

        audio_io_frame = QFrame()
        audio_io_frame.setObjectName("ioFrame")
        audio_io_layout = QVBoxLayout(audio_io_frame)
        
        self.audio_input_label = QLabel(self.lang["input_audio_label"])
        audio_io_layout.addWidget(self.audio_input_label)
        audio_input_layout = QHBoxLayout()
        self.audio_input_path_edit = QLineEdit()
        self.audio_input_path_edit.setReadOnly(True)
        self.audio_browse_btn = QPushButton(self.lang["browse_btn"])
        self.audio_browse_btn.clicked.connect(self.browse_audio_input_file)
        audio_input_layout.addWidget(self.audio_input_path_edit)
        audio_input_layout.addWidget(self.audio_browse_btn)
        audio_io_layout.addLayout(audio_input_layout)

        self.audio_output_label = QLabel(self.lang["output_audio_label"])
        audio_io_layout.addWidget(self.audio_output_label)
        audio_output_layout = QHBoxLayout()
        self.audio_output_path_edit = QLineEdit()
        self.audio_output_path_edit.setPlaceholderText(self.lang["output_placeholder_audio"])
        self.audio_output_path_edit.setReadOnly(True)
        self.audio_browse_output_btn = QPushButton(self.lang["browse_btn"])
        self.audio_browse_output_btn.clicked.connect(self.browse_audio_output_folder)
        audio_output_layout.addWidget(self.audio_output_path_edit)
        audio_output_layout.addWidget(self.audio_browse_output_btn)
        audio_io_layout.addLayout(audio_output_layout)

        self.audio_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.audio_batch_mode_checkbox.stateChanged.connect(self._update_audio_input_ui)
        audio_io_layout.addWidget(self.audio_batch_mode_checkbox)
        layout.addWidget(audio_io_frame)

        audio_settings_frame = QFrame()
        audio_settings_frame.setObjectName("ioFrame")
        audio_settings_layout = QVBoxLayout(audio_settings_frame)
        self.audio_options_label = QLabel(self.lang["audio_options_label"])
        audio_settings_layout.addWidget(self.audio_options_label)
        settings_hbox = QHBoxLayout()
        
        self.audio_format_label = QLabel(self.lang["output_format_label"])
        settings_hbox.addWidget(self.audio_format_label)
        self.audio_format_combo = QComboBox()
        self.audio_format_combo.addItems(self.lang["audio_formats"])
        settings_hbox.addWidget(self.audio_format_combo, 1)
        settings_hbox.addSpacing(20)
        self.audio_bitrate_label = QLabel(self.lang["bitrate_label"])
        settings_hbox.addWidget(self.audio_bitrate_label)
        self.audio_bitrate_combo = QComboBox()
        self.audio_bitrate_combo.addItems(["96 kbps", "128 kbps", "192 kbps", "256 kbps", "320 kbps"])
        self.audio_bitrate_combo.setCurrentText("192 kbps")
        settings_hbox.addWidget(self.audio_bitrate_combo, 1)
        audio_settings_layout.addLayout(settings_hbox)
        layout.addWidget(audio_settings_frame)

        self.audio_convert_btn = QPushButton(self.lang["start_audio_conv_btn"])
        self.audio_convert_btn.setObjectName("convertButton")
        self.audio_convert_btn.clicked.connect(self.start_audio_conversion)
        layout.addWidget(self.audio_convert_btn)

        self.audio_progress_bar = QProgressBar()
        self.audio_progress_bar.setValue(0)
        self.audio_progress_bar.setTextVisible(False)
        layout.addWidget(self.audio_progress_bar)
        
        self.audio_status_label = QLabel(self.lang["ready_status_audio"])
        self.audio_status_label.setObjectName("statusLabel")
        layout.addWidget(self.audio_status_label)

        layout.addStretch()
        self._update_audio_input_ui()
        return tab_widget

    # [MODIFIKASI] Tab baru untuk konversi Video ke Audio
    def _create_video_audio_converter_tab(self):
        tab_widget = QWidget()
        layout = QVBoxLayout(tab_widget)
        layout.setContentsMargins(10, 20, 10, 10)
        layout.setSpacing(15)

        va_io_frame = QFrame()
        va_io_frame.setObjectName("ioFrame")
        va_io_layout = QVBoxLayout(va_io_frame)
        
        self.va_input_label = QLabel(self.lang["input_va_label"])
        va_io_layout.addWidget(self.va_input_label)
        va_input_layout = QHBoxLayout()
        self.va_input_path_edit = QLineEdit()
        self.va_input_path_edit.setReadOnly(True)
        self.va_browse_btn = QPushButton(self.lang["browse_btn"])
        self.va_browse_btn.clicked.connect(self.browse_va_input_file)
        va_input_layout.addWidget(self.va_input_path_edit)
        va_input_layout.addWidget(self.va_browse_btn)
        va_io_layout.addLayout(va_input_layout)

        self.va_output_label = QLabel(self.lang["output_va_label"])
        va_io_layout.addWidget(self.va_output_label)
        va_output_layout = QHBoxLayout()
        self.va_output_path_edit = QLineEdit()
        self.va_output_path_edit.setPlaceholderText(self.lang["output_placeholder_va"])
        self.va_output_path_edit.setReadOnly(True)
        self.va_browse_output_btn = QPushButton(self.lang["browse_btn"])
        self.va_browse_output_btn.clicked.connect(self.browse_va_output_folder)
        va_output_layout.addWidget(self.va_output_path_edit)
        va_output_layout.addWidget(self.va_browse_output_btn)
        va_io_layout.addLayout(va_output_layout)

        self.va_batch_mode_checkbox = QCheckBox(self.lang["batch_mode_checkbox"])
        self.va_batch_mode_checkbox.stateChanged.connect(self._update_va_input_ui)
        va_io_layout.addWidget(self.va_batch_mode_checkbox)
        layout.addWidget(va_io_frame)

        va_settings_frame = QFrame()
        va_settings_frame.setObjectName("ioFrame")
        va_settings_layout = QVBoxLayout(va_settings_frame)
        self.va_options_label = QLabel(self.lang["va_options_label"])
        va_settings_layout.addWidget(self.va_options_label)
        settings_hbox = QHBoxLayout()
        
        self.va_format_label = QLabel(self.lang["output_format_label"])
        settings_hbox.addWidget(self.va_format_label)
        self.va_format_combo = QComboBox()
        self.va_format_combo.addItems(self.lang["audio_formats"])
        settings_hbox.addWidget(self.va_format_combo, 1)
        settings_hbox.addSpacing(20)
        self.va_bitrate_label = QLabel(self.lang["bitrate_label"])
        settings_hbox.addWidget(self.va_bitrate_label)
        self.va_bitrate_combo = QComboBox()
        self.va_bitrate_combo.addItems(["96 kbps", "128 kbps", "192 kbps", "256 kbps", "320 kbps"])
        self.va_bitrate_combo.setCurrentText("192 kbps")
        settings_hbox.addWidget(self.va_bitrate_combo, 1)
        va_settings_layout.addLayout(settings_hbox)
        layout.addWidget(va_settings_frame)

        self.va_convert_btn = QPushButton(self.lang["start_va_conv_btn"])
        self.va_convert_btn.setObjectName("convertButton")
        self.va_convert_btn.clicked.connect(self.start_va_conversion)
        layout.addWidget(self.va_convert_btn)

        self.va_progress_bar = QProgressBar()
        self.va_progress_bar.setValue(0)
        self.va_progress_bar.setTextVisible(False)
        layout.addWidget(self.va_progress_bar)
        
        self.va_status_label = QLabel(self.lang["ready_status_va"])
        self.va_status_label.setObjectName("statusLabel")
        layout.addWidget(self.va_status_label)

        layout.addStretch()
        self._update_va_input_ui()
        return tab_widget

    def _apply_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow, QWidget {
                background-color: #2E3440; color: #ECEFF4; font-family: Segoe UI, sans-serif;
            }
            QTabWidget::pane { border-top: 2px solid #434C5E; }
            QTabBar::tab {
                background: #3B4252; color: #D8DEE9; padding: 10px;
                border: 1px solid #434C5E; border-bottom: none; border-top-left-radius: 4px; border-top-right-radius: 4px;
            }
            QTabBar::tab:selected { background: #434C5E; color: #ECEFF4; }
            #titleLabel { font-size: 24pt; font-weight: bold; color: #88C0D0; padding-bottom: 10px; }
            #statusLabel { color: #A3BE8C; }
            QLabel { font-size: 10pt; }
            QCheckBox { font-size: 9pt; padding-top: 5px; }
            QLineEdit {
                background-color: #4C566A; border: 1px solid #5E81AC; border-radius: 4px;
                padding: 6px; color: #D8DEE9;
            }
            QPushButton {
                background-color: #5E81AC; color: #ECEFF4; border: none; padding: 8px 16px;
                border-radius: 4px; font-size: 10pt;
            }
            QPushButton:hover { background-color: #81A1C1; }
            QPushButton:pressed { background-color: #4C566A; }
            #convertButton { background-color: #A3BE8C; font-weight: bold; color: #2E3440; }
            #convertButton:hover { background-color: #B48EAD; }
            QComboBox {
                background-color: #4C566A; border: 1px solid #5E81AC; border-radius: 4px; padding: 6px;
            }
            QComboBox::drop-down { border: none; }
            QProgressBar {
                border: 1px solid #4C566A; border-radius: 4px; text-align: center; height: 10px;
            }
            QProgressBar::chunk { background-color: #88C0D0; border-radius: 4px; }
            #ioFrame, #icoFrame {
                border: 1px solid #434C5E; border-radius: 5px; padding: 10px;
            }
            #aboutText {
                background-color: #3B4252;
                border: 1px solid #434C5E;
                color: #ECEFF4;
                font-size: 11pt;
                padding: 10px;
            }
        """)

    # --- [MODIFIKASI] Logika untuk File Converter ---
    def browse_input_file(self):
        mode = self.mode_combo.currentText()
        
        # Daftar filter file berdasarkan mode
        filters = {
            self.lang["modes"][0]: "PDF Files (*.pdf)",           # PDF ke Gambar
            self.lang["modes"][1]: "PNG Files (*.png)",           # PNG ke ICO
            self.lang["modes"][2]: "PDF Files (*.pdf)",           # PDF ke TXT
            self.lang["modes"][3]: "PDF Files (*.pdf)",           # PDF ke DOCX
            self.lang["modes"][4]: "PDF Files (*.pdf)",           # PDF ke XLSX
            self.lang["modes"][5]: "Word Documents (*.docx)",     # DOCX ke PDF
            self.lang["modes"][6]: "Image Files (*.png *.jpg *.jpeg *.bmp *.webp *.gif)", # Gambar ke PDF
        }
        filter_str = filters.get(mode, "All Files (*)")
        
        # Mode Gambar ke PDF memungkinkan multiple file
        is_multi_file_mode = (mode == self.lang["modes"][6])

        if is_multi_file_mode:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_label"], "", filter_str)
            if files:
                self.file_batch_files = files
                self.input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            self.file_batch_files = []
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_label"], "", filter_str)
            if file_path:
                self.input_path_edit.setText(file_path)

    def browse_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_label"])
        if folder_path:
            self.output_path_edit.setText(folder_path)

    def _update_ui_for_mode(self, mode):
        # Tentukan opsi mana yang akan ditampilkan untuk setiap mode
        is_pdf_to_image = mode == self.lang["modes"][0]
        is_png_to_ico = mode == self.lang["modes"][1]
        is_image_to_pdf = mode == self.lang["modes"][6]

        self.format_combo.setVisible(is_pdf_to_image)
        self.format_label.setVisible(is_pdf_to_image)
        self.ico_sizes_frame.setVisible(is_png_to_ico)
        
        # Atur ulang input field
        self.input_path_edit.clear()
        self.file_batch_files = []
        if is_image_to_pdf:
            self.input_path_edit.setPlaceholderText("Pilih satu atau lebih file gambar...")
        else:
            self.input_path_edit.setPlaceholderText(self.lang["input_placeholder_file"])
            
        self.progress_bar.setValue(0)
        self.status_label.setText(self.lang["ready_status"])

    def start_file_conversion(self):
        output_path = self.output_path_edit.text()
        mode = self.mode_combo.currentText()
        out_format = self.format_combo.currentText()
        is_multi_file_mode = (mode == self.lang["modes"][6])
        
        input_path = self.file_batch_files if is_multi_file_mode else self.input_path_edit.text()

        # Validasi input
        if not input_path or (not is_multi_file_mode and not os.path.exists(input_path)):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"])
            return
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"])
            return

        ico_sizes = []
        if mode == self.lang["modes"][1]: # PNG to ICO
            for size_str, checkbox in self.ico_checkboxes.items():
                if checkbox.isChecked():
                    w, h = map(int, size_str.split('x'))
                    ico_sizes.append((w, h))
            if not ico_sizes:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["no_ico_size_msg"])
                return

        self.convert_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        
        self.thread = QThread()
        self.worker = FileConversionWorker(input_path, output_path, mode, out_format, ico_sizes, self.lang)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(lambda v, t: self.update_progress(self.progress_bar, self.status_label, v, t))
        self.worker.conversion_finished.connect(self.on_file_conversion_finished)
        self.worker.conversion_error.connect(lambda msg: self.conversion_error(self.convert_btn, self.progress_bar, self.status_label, msg))
        
        self.thread.start()
        
    def on_file_conversion_finished(self, msg):
        self.conversion_finished(self.convert_btn, self.progress_bar, self.status_label, msg)

    # --- Logika untuk Image Converter ---
    def _update_image_input_ui(self):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        if is_batch:
            self.img_input_path_edit.setPlaceholderText(self.lang["input_placeholder_multi_img"])
        else:
            self.img_input_path_edit.setPlaceholderText(self.lang["input_placeholder_single_img"])
        self.img_input_path_edit.clear()
        self.image_batch_files = []

    def _update_image_options(self):
        selected_format = self.img_format_combo.currentText().lower()
        has_quality = selected_format in ['jpeg', 'jpg', 'webp']
        self.img_quality_label.setVisible(has_quality)
        self.img_quality_combo.setVisible(has_quality)

    def browse_image_input_file(self):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        filter_str = "Image Files (*.png *.jpg *.jpeg *.bmp *.webp *.gif)"
        
        if is_batch:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_img_label"], "", filter_str)
            if files:
                self.image_batch_files = files
                self.img_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_img_label"], "", filter_str)
            if file_path:
                self.img_input_path_edit.setText(file_path)

    def browse_image_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_img_label"])
        if folder_path:
            self.img_output_path_edit.setText(folder_path)

    def start_image_conversion(self):
        output_path = self.img_output_path_edit.text()
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"])
            return

        is_batch = self.img_batch_mode_checkbox.isChecked()
        if is_batch:
            if not self.image_batch_files:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["batch_no_files_msg"])
                return
            self.current_image_batch_index = 0
            self._start_next_batch_image_conversion()
        else:
            input_path = self.img_input_path_edit.text()
            if not input_path or not os.path.exists(input_path):
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"])
                return
            self._start_single_image_conversion(input_path)

    def _start_single_image_conversion(self, input_path):
        self.img_convert_btn.setEnabled(False)
        self.img_progress_bar.setValue(0)
        self.img_status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))

        out_format = self.img_format_combo.currentText()
        resolution = self.img_resolution_combo.currentText()
        quality = self.img_quality_combo.currentText()
        output_path = self.img_output_path_edit.text()

        self.thread = QThread()
        self.worker = ImageConversionWorker(input_path, output_path, out_format, resolution, quality, self.lang)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self.on_image_progress_update)
        self.worker.conversion_finished.connect(self.on_image_conversion_finished)
        self.worker.conversion_error.connect(self.on_image_conversion_error)
        
        self.thread.start()

    def _start_next_batch_image_conversion(self):
        if self.current_image_batch_index < len(self.image_batch_files):
            input_path = self.image_batch_files[self.current_image_batch_index]
            self.img_status_label.setText(
                self.lang["converting_batch_file"].format(
                    current=self.current_image_batch_index + 1,
                    total=len(self.image_batch_files),
                    filename=os.path.basename(input_path)
                )
            )
            self._start_single_image_conversion(input_path)
        else:
            final_msg = self.lang["batch_complete_msg"].format(count=len(self.image_batch_files))
            self.img_status_label.setText(final_msg)
            self.img_convert_btn.setEnabled(True)
            QMessageBox.information(self, self.lang["batch_complete_title"], final_msg)
            self._open_output_folder(self.img_output_path_edit.text())
    
    def on_image_progress_update(self, value, text):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        if is_batch and self.image_batch_files:
            current_file_info = f"File {self.current_image_batch_index + 1}/{len(self.image_batch_files)}"
            self.update_progress(self.img_progress_bar, self.img_status_label, value, f"({current_file_info}) {text}")
        else:
            self.update_progress(self.img_progress_bar, self.img_status_label, value, text)

    def on_image_conversion_finished(self, msg):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        if is_batch and self.image_batch_files:
            self.img_progress_bar.setValue(100)
            self.current_image_batch_index += 1
            if self.thread:
                self.thread.quit()
                self.thread.wait()
            self._start_next_batch_image_conversion()
        else:
            self.conversion_finished(self.img_convert_btn, self.img_progress_bar, self.img_status_label, msg)

    def on_image_conversion_error(self, msg):
        is_batch = self.img_batch_mode_checkbox.isChecked()
        if is_batch and self.image_batch_files:
            error_msg = self.lang["error_in_batch"].format(index=self.current_image_batch_index + 1, error=msg)
            self.conversion_error(self.img_convert_btn, self.img_progress_bar, self.img_status_label, error_msg)
        else:
            self.conversion_error(self.img_convert_btn, self.img_progress_bar, self.img_status_label, msg)

    # --- Logika untuk Video Converter ---
    def _update_video_input_ui(self):
        is_batch = self.batch_mode_checkbox.isChecked()
        if is_batch:
            self.vid_input_path_edit.setPlaceholderText(self.lang["input_placeholder_multi_vid"])
        else:
            self.vid_input_path_edit.setPlaceholderText(self.lang["input_placeholder_single_vid"])
        self.vid_input_path_edit.clear()
        self.video_batch_files = []

    def browse_video_input_file(self):
        is_batch = self.batch_mode_checkbox.isChecked()
        filter_str = "Video Files (*.mp4 *.mkv *.avi *.mov *.flv *.wmv)"
        
        if is_batch:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_vid_label"], "", filter_str)
            if files:
                self.video_batch_files = files
                self.vid_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_vid_label"], "", filter_str)
            if file_path:
                self.vid_input_path_edit.setText(file_path)

    def browse_video_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_vid_label"])
        if folder_path:
            self.vid_output_path_edit.setText(folder_path)
            
    def _find_ffmpeg(self):
        local_path = "ffmpeg.exe" if os.name == 'nt' else "ffmpeg"
        base_path = sys._MEIPASS if hasattr(sys, "_MEIPASS") else os.path.dirname(os.path.abspath(__file__))
        local_ffmpeg = os.path.join(base_path, local_path)
        if os.path.exists(local_ffmpeg): return local_ffmpeg
        from shutil import which
        if which("ffmpeg"): return "ffmpeg"
        return None

    def start_video_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path:
            QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"])
            return

        output_path = self.vid_output_path_edit.text()
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"])
            return

        is_batch = self.batch_mode_checkbox.isChecked()
        if is_batch:
            if not self.video_batch_files:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["batch_no_files_msg"])
                return
            self.current_batch_index = 0
            self._start_next_batch_video_conversion()
        else:
            input_path = self.vid_input_path_edit.text()
            if not input_path or not os.path.exists(input_path):
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"])
                return
            self._start_single_video_conversion(input_path)
    
    def _start_single_video_conversion(self, input_path):
        self.vid_convert_btn.setEnabled(False)
        self.vid_progress_bar.setValue(0)
        self.vid_status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))

        out_format = self.vid_format_combo.currentText()
        resolution = self.vid_resolution_combo.currentText()
        quality = self.vid_quality_combo.currentText()
        output_path = self.vid_output_path_edit.text()

        self.thread = QThread()
        self.worker = VideoConversionWorker(self.ffmpeg_path, input_path, output_path, out_format, resolution, quality, self.lang)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self.on_video_progress_update)
        self.worker.conversion_finished.connect(self.on_video_conversion_finished)
        self.worker.conversion_error.connect(self.on_video_conversion_error)
        
        self.thread.start()
        
    def _start_next_batch_video_conversion(self):
        if self.current_batch_index < len(self.video_batch_files):
            input_path = self.video_batch_files[self.current_batch_index]
            self.vid_status_label.setText(
                self.lang["converting_batch_file"].format(
                    current=self.current_batch_index + 1,
                    total=len(self.video_batch_files),
                    filename=os.path.basename(input_path)
                )
            )
            self._start_single_video_conversion(input_path)
        else:
            final_msg = self.lang["batch_complete_msg"].format(count=len(self.video_batch_files))
            self.vid_status_label.setText(final_msg)
            self.vid_convert_btn.setEnabled(True)
            QMessageBox.information(self, self.lang["batch_complete_title"], final_msg)
            self._open_output_folder(self.vid_output_path_edit.text())
            self._handle_shutdown() # [MODIFIKASI]

    def on_video_progress_update(self, value, text):
        is_batch = self.batch_mode_checkbox.isChecked()
        if is_batch and self.video_batch_files:
            current_file_info = f"File {self.current_batch_index + 1}/{len(self.video_batch_files)}"
            self.update_progress(self.vid_progress_bar, self.vid_status_label, value, f"({current_file_info}) {text}")
        else:
            self.update_progress(self.vid_progress_bar, self.vid_status_label, value, text)

    def on_video_conversion_finished(self, msg):
        is_batch = self.batch_mode_checkbox.isChecked()
        if is_batch and self.video_batch_files:
            self.vid_progress_bar.setValue(100)
            self.current_batch_index += 1
            if self.thread:
                self.thread.quit()
                self.thread.wait()
            self._start_next_batch_video_conversion()
        else:
            self.conversion_finished(self.vid_convert_btn, self.vid_progress_bar, self.vid_status_label, msg)
            self._handle_shutdown() # [MODIFIKASI]

    def on_video_conversion_error(self, msg):
        is_batch = self.batch_mode_checkbox.isChecked()
        if is_batch and self.video_batch_files:
            error_msg = self.lang["error_in_batch"].format(index=self.current_batch_index + 1, error=msg)
            self.conversion_error(self.vid_convert_btn, self.vid_progress_bar, self.vid_status_label, error_msg)
        else:
            self.conversion_error(self.vid_convert_btn, self.vid_progress_bar, self.vid_status_label, msg)

    # --- Logika untuk Audio Converter ---
    def _update_audio_input_ui(self):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        if is_batch:
            self.audio_input_path_edit.setPlaceholderText(self.lang["input_placeholder_multi_audio"])
        else:
            self.audio_input_path_edit.setPlaceholderText(self.lang["input_placeholder_single_audio"])
        self.audio_input_path_edit.clear()
        self.audio_batch_files = []

    def browse_audio_input_file(self):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        filter_str = "Audio Files (*.mp3 *.wav *.aac *.flac *.ogg *.wma *.m4a)"
        
        if is_batch:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_audio_label"], "", filter_str)
            if files:
                self.audio_batch_files = files
                self.audio_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_audio_label"], "", filter_str)
            if file_path:
                self.audio_input_path_edit.setText(file_path)

    def browse_audio_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_audio_label"])
        if folder_path:
            self.audio_output_path_edit.setText(folder_path)

    def start_audio_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path:
            QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"])
            return

        output_path = self.audio_output_path_edit.text()
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"])
            return

        is_batch = self.audio_batch_mode_checkbox.isChecked()
        if is_batch:
            if not self.audio_batch_files:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["batch_no_files_msg"])
                return
            self.current_audio_batch_index = 0
            self._start_next_batch_audio_conversion()
        else:
            input_path = self.audio_input_path_edit.text()
            if not input_path or not os.path.exists(input_path):
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"])
                return
            self._start_single_audio_conversion(input_path)

    def _start_single_audio_conversion(self, input_path):
        self.audio_convert_btn.setEnabled(False)
        self.audio_progress_bar.setValue(0)
        self.audio_status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))

        out_format = self.audio_format_combo.currentText()
        bitrate = self.audio_bitrate_combo.currentText()
        output_path = self.audio_output_path_edit.text()

        self.thread = QThread()
        self.worker = AudioConversionWorker(self.ffmpeg_path, input_path, output_path, out_format, bitrate, self.lang)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self.on_audio_progress_update)
        self.worker.conversion_finished.connect(self.on_audio_conversion_finished)
        self.worker.conversion_error.connect(self.on_audio_conversion_error)
        
        self.thread.start()

    def _start_next_batch_audio_conversion(self):
        if self.current_audio_batch_index < len(self.audio_batch_files):
            input_path = self.audio_batch_files[self.current_audio_batch_index]
            self.audio_status_label.setText(
                self.lang["converting_batch_file"].format(
                    current=self.current_audio_batch_index + 1,
                    total=len(self.audio_batch_files),
                    filename=os.path.basename(input_path)
                )
            )
            self._start_single_audio_conversion(input_path)
        else:
            final_msg = self.lang["batch_complete_msg"].format(count=len(self.audio_batch_files))
            self.audio_status_label.setText(final_msg)
            self.audio_convert_btn.setEnabled(True)
            QMessageBox.information(self, self.lang["batch_complete_title"], final_msg)
            self._open_output_folder(self.audio_output_path_edit.text())

    def on_audio_progress_update(self, value, text):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        if is_batch and self.audio_batch_files:
            current_file_info = f"File {self.current_audio_batch_index + 1}/{len(self.audio_batch_files)}"
            self.update_progress(self.audio_progress_bar, self.audio_status_label, value, f"({current_file_info}) {text}")
        else:
            self.update_progress(self.audio_progress_bar, self.audio_status_label, value, text)

    def on_audio_conversion_finished(self, msg):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        if is_batch and self.audio_batch_files:
            self.audio_progress_bar.setValue(100)
            self.current_audio_batch_index += 1
            if self.thread:
                self.thread.quit()
                self.thread.wait()
            self._start_next_batch_audio_conversion()
        else:
            self.conversion_finished(self.audio_convert_btn, self.audio_progress_bar, self.audio_status_label, msg)

    def on_audio_conversion_error(self, msg):
        is_batch = self.audio_batch_mode_checkbox.isChecked()
        if is_batch and self.audio_batch_files:
            error_msg = self.lang["error_in_batch"].format(index=self.current_audio_batch_index + 1, error=msg)
            self.conversion_error(self.audio_convert_btn, self.audio_progress_bar, self.audio_status_label, error_msg)
        else:
            self.conversion_error(self.audio_convert_btn, self.audio_progress_bar, self.audio_status_label, msg)
            
    # --- [MODIFIKASI] Logika untuk Video-Audio Converter ---
    def _update_va_input_ui(self):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        if is_batch:
            self.va_input_path_edit.setPlaceholderText(self.lang["input_placeholder_multi_va"])
        else:
            self.va_input_path_edit.setPlaceholderText(self.lang["input_placeholder_single_va"])
        self.va_input_path_edit.clear()
        self.va_batch_files = []

    def browse_va_input_file(self):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        filter_str = "Video Files (*.mp4 *.mkv *.avi *.mov *.flv *.wmv)"
        
        if is_batch:
            files, _ = QFileDialog.getOpenFileNames(self, self.lang["input_va_label"], "", filter_str)
            if files:
                self.va_batch_files = files
                self.va_input_path_edit.setText(f"{len(files)} file(s) selected.")
        else:
            file_path, _ = QFileDialog.getOpenFileName(self, self.lang["input_va_label"], "", filter_str)
            if file_path:
                self.va_input_path_edit.setText(file_path)

    def browse_va_output_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.lang["output_va_label"])
        if folder_path:
            self.va_output_path_edit.setText(folder_path)

    def start_va_conversion(self):
        self.ffmpeg_path = self._find_ffmpeg()
        if not self.ffmpeg_path:
            QMessageBox.critical(self, self.lang["ffmpeg_not_found_title"], self.lang["ffmpeg_not_found_msg"])
            return

        output_path = self.va_output_path_edit.text()
        if not output_path or not os.path.isdir(output_path):
            QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_output_folder_msg"])
            return

        is_batch = self.va_batch_mode_checkbox.isChecked()
        if is_batch:
            if not self.va_batch_files:
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["batch_no_files_msg"])
                return
            self.current_va_batch_index = 0
            self._start_next_batch_va_conversion()
        else:
            input_path = self.va_input_path_edit.text()
            if not input_path or not os.path.exists(input_path):
                QMessageBox.warning(self, self.lang["invalid_input_title"], self.lang["invalid_input_file_msg"])
                return
            self._start_single_va_conversion(input_path)

    def _start_single_va_conversion(self, input_path):
        self.va_convert_btn.setEnabled(False)
        self.va_progress_bar.setValue(0)
        self.va_status_label.setText(self.lang["preparing_conversion"].format(filename=os.path.basename(input_path)))

        out_format = self.va_format_combo.currentText()
        bitrate = self.va_bitrate_combo.currentText()
        output_path = self.va_output_path_edit.text()

        self.thread = QThread()
        # Menggunakan kembali AudioConversionWorker karena logikanya sama
        self.worker = AudioConversionWorker(self.ffmpeg_path, input_path, output_path, out_format, bitrate, self.lang)
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.progress_updated.connect(self.on_va_progress_update)
        self.worker.conversion_finished.connect(self.on_va_conversion_finished)
        self.worker.conversion_error.connect(self.on_va_conversion_error)
        
        self.thread.start()

    def _start_next_batch_va_conversion(self):
        if self.current_va_batch_index < len(self.va_batch_files):
            input_path = self.va_batch_files[self.current_va_batch_index]
            self.va_status_label.setText(
                self.lang["converting_batch_file"].format(
                    current=self.current_va_batch_index + 1,
                    total=len(self.va_batch_files),
                    filename=os.path.basename(input_path)
                )
            )
            self._start_single_va_conversion(input_path)
        else:
            final_msg = self.lang["batch_complete_msg"].format(count=len(self.va_batch_files))
            self.va_status_label.setText(final_msg)
            self.va_convert_btn.setEnabled(True)
            QMessageBox.information(self, self.lang["batch_complete_title"], final_msg)
            self._open_output_folder(self.va_output_path_edit.text())

    def on_va_progress_update(self, value, text):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        if is_batch and self.va_batch_files:
            current_file_info = f"File {self.current_va_batch_index + 1}/{len(self.va_batch_files)}"
            self.update_progress(self.va_progress_bar, self.va_status_label, value, f"({current_file_info}) {text}")
        else:
            self.update_progress(self.va_progress_bar, self.va_status_label, value, text)

    def on_va_conversion_finished(self, msg):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        if is_batch and self.va_batch_files:
            self.va_progress_bar.setValue(100)
            self.current_va_batch_index += 1
            if self.thread:
                self.thread.quit()
                self.thread.wait()
            self._start_next_batch_va_conversion()
        else:
            self.conversion_finished(self.va_convert_btn, self.va_progress_bar, self.va_status_label, msg)

    def on_va_conversion_error(self, msg):
        is_batch = self.va_batch_mode_checkbox.isChecked()
        if is_batch and self.va_batch_files:
            error_msg = self.lang["error_in_batch"].format(index=self.current_va_batch_index + 1, error=msg)
            self.conversion_error(self.va_convert_btn, self.va_progress_bar, self.va_status_label, error_msg)
        else:
            self.conversion_error(self.va_convert_btn, self.va_progress_bar, self.va_status_label, msg)


    # --- Slot Generik dan Fungsi Bantuan ---
    def update_progress(self, progress_bar, status_label, value, text):
        progress_bar.setValue(value)
        status_label.setText(text)
    
    # [MODIFIKASI] Fungsi untuk mematikan PC
    def _handle_shutdown(self):
        if self.shutdown_checkbox.isChecked():
            reply = QMessageBox.question(self, self.lang["shutdown_confirm_title"],
                                         self.lang["shutdown_confirm_msg"],
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                if os.name == 'nt': # Windows
                    os.system("shutdown /s /t 30") # 30 detik countdown
                else: # Linux atau MacOS
                    # Peringatan: Perintah ini mungkin memerlukan hak akses sudo.
                    os.system("shutdown -h +1") # Matikan dalam 1 menit

    def _open_output_folder(self, path):
        if not os.path.isdir(path): return
        try:
            if sys.platform == "win32":
                os.startfile(path)
            elif sys.platform == "darwin": # macOS
                subprocess.Popen(["open", path])
            else: # Linux
                subprocess.Popen(["xdg-open", path])
        except Exception as e:
            print(f"Gagal membuka folder output: {e}")

    def conversion_finished(self, button, progress_bar, status_label, message):
        status_label.setText(message)
        progress_bar.setValue(100)
        button.setEnabled(True)
        if self.thread:
            self.thread.quit()
            self.thread.wait()
        
        QMessageBox.information(self, self.lang["done"], message)
        
        output_path = ""
        # Mencari tahu path output dari tab mana yang aktif
        current_tab_index = self.tab_widget.currentIndex()
        if current_tab_index == 0: output_path = self.output_path_edit.text()
        elif current_tab_index == 1: output_path = self.img_output_path_edit.text()
        elif current_tab_index == 2: output_path = self.vid_output_path_edit.text()
        elif current_tab_index == 3: output_path = self.audio_output_path_edit.text()
        elif current_tab_index == 4: output_path = self.va_output_path_edit.text()

        if output_path:
            self._open_output_folder(output_path)

    def conversion_error(self, button, progress_bar, status_label, message):
        status_label.setText(f"{self.lang['error_title']}: {message}")
        progress_bar.setValue(0)
        button.setEnabled(True)
        if self.thread:
            self.thread.quit()
            self.thread.wait()
        QMessageBox.critical(self, self.lang["error_title"], message)

    def closeEvent(self, event):
        if self.thread and self.thread.isRunning():
            self.worker.stop()
            self.thread.quit()
            self.thread.wait()
        event.accept()

# --- Entry Point Aplikasi ---
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ConverterApp()
    window.show()
    sys.exit(app.exec())